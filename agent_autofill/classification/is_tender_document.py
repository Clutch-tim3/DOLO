"""
Is this a tender document?

One Haiku call, on the first ~1500 characters, returning a confidence and a
one-line reason. Everything about this module is shaped by the fact that it is
the only part of Agent Autofill that costs money per document, and that it
runs over whatever happens to be in a user's connected folder.

Order of refusals, cheapest first:

  1. Extension / magic bytes  — PDF and DOCX only. A .jpg never reaches Haiku.
  2. Empty or unreadable text — a scanned PDF with no text layer has nothing
                                to classify; say so rather than guessing.
  3. Global rate limit        — agent/rate_limiter.py, shared with every other
                                Anthropic call in the app. No parallel limiter.
  4. The model call.

Threshold: confidence >= 0.7 proceeds. Below that the document is logged and
skipped. The asymmetry is deliberate — a missed tender costs the user one
manual upload, while a false positive spends the rest of the pipeline (and a
slice of their daily quota) on a document that was never a tender.

The tier/quota gate is NOT here. It lives in agent/subscription.py and is the
orchestrator's first step, because it must run before this module is even
imported into the request path.
"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

# Imported as a module, not as a bare name, so a test can monkeypatch
# `claude_client.call_claude_with_tracking` and actually intercept the call.
# `from ... import call_claude_with_tracking` would bind a copy and make the
# "no request was made" proof unverifiable.
from agent import claude_client
from agent import rate_limiter
from agent_autofill.extraction.legacy_doc_reader import detect_format

log = logging.getLogger("agent_autofill.classification")

#: Haiku 4.5. `claude-3-5-haiku-20241022` was retired on 2026-02-19 and 404s —
#: do not "restore" it. The dated alias claude-haiku-4-5-20251001 is the same
#: model at the same price if a pin is ever needed.
CLASSIFIER_MODEL = "claude-haiku-4-5"

#: Proceed at or above this. Chosen before looking at any scores; see the
#: module docstring for why it is not symmetric.
CONFIDENCE_THRESHOLD = 0.7

#: How much of the document the model sees. The identifying material of a
#: tender — issuing organ of state, bid number, closing date, "REQUEST FOR
#: QUOTATION" — is in the first page, and paying for 145 pages to learn that
#: would defeat the purpose of the gate.
HEAD_CHARS = 1500

#: Nothing else is classified. Legacy .doc is detected separately and reported
#: as read-only by the orchestrator, not silently dropped here.
SUPPORTED_SUFFIXES = {".pdf", ".docx"}

#: Below this there is not enough text to judge — usually a scanned page with
#: no text layer.
MIN_TEXT_CHARS = 40

MAX_TOKENS = 300

SYSTEM_PROMPT = (
    "You classify documents for a South African procurement assistant. You are "
    "given the opening text of one file taken from a supplier's document folder. "
    "Decide whether it is a government or institutional procurement document that "
    "a supplier would fill in and submit: a tender, bid, RFQ, RFP, RFI, RFB, "
    "expression of interest, or a bid form pack (SBD/MBD forms, pricing schedules, "
    "returnable schedules). South African issuers are the common case (municipalities, "
    "provincial and national departments, SOEs such as Eskom or Transnet, universities), "
    "but international government or institutional tenders count too.\n\n"
    "These are NOT tender documents, even when they mention a tender: company "
    "registration certificates (CIPC / COR forms), SARS tax clearance or tax "
    "compliance pins, B-BBEE certificates, CSD supplier reports, bank confirmation "
    "letters, invoices, quotations the supplier issued, ID documents, CVs, "
    "insurance policies, and award or regret letters after the fact.\n\n"
    "Reply with a single JSON object and nothing else:\n"
    '{"is_tender": true|false, "confidence": 0.0-1.0, "document_type": "<short label>", '
    '"reason": "<one short sentence>"}\n\n'
    "confidence is your confidence in the is_tender verdict you gave. If the text is "
    "too short, garbled or ambiguous to judge, say is_tender false with a low "
    "confidence and say so in the reason."
)


@dataclass
class ClassificationResult:
    """What the gate decided, and what it cost."""

    path: str
    is_tender: bool = False
    confidence: float = 0.0
    reason: str = ""
    document_type: str = ""
    #: "classified" | "unsupported_type" | "unreadable" | "rate_limited" | "error"
    status: str = "classified"
    #: True only when the model was actually invoked.
    api_called: bool = False
    text_head: str = field(default="", repr=False)
    #: True when the text above came from OCR rather than off the page. Worth
    #: carrying all the way out: a confidence built on recognised text deserves
    #: to be read differently from one built on the document's own words.
    ocr_used: bool = False
    #: What to tell the user about the scan — that OCR read it, that OCR is not
    #: configured, or that it ran and found nothing. Empty when OCR never ran.
    ocr_note: str = ""

    @property
    def proceed(self) -> bool:
        """The single question the orchestrator asks."""
        return (
            self.status == "classified"
            and self.is_tender
            and self.confidence >= CONFIDENCE_THRESHOLD
        )

    @property
    def summary_line(self) -> str:
        return (f"{Path(self.path).name}: {'TENDER' if self.is_tender else 'not a tender'} "
                f"@ {self.confidence:.2f} — {self.reason}")


def is_classifiable(path: str | Path) -> tuple[bool, str]:
    """
    Cheap pre-flight. Returns (ok, reason).

    Checks the extension first (free) and then the magic bytes (one 8-byte
    read), because a .doc renamed to .docx is common enough that trusting the
    name here would send an unreadable file to the model.
    """
    path = Path(path)
    if path.suffix.lower() not in SUPPORTED_SUFFIXES:
        return False, (f"{path.suffix or 'no extension'} is not a supported document type — "
                       f"Agent Autofill reads PDF and DOCX only.")
    if not path.exists():
        return False, f"File not found: {path.name}"

    fmt = detect_format(path)
    if fmt == "doc":
        return False, (f"{path.name} is a legacy Word .doc file despite its name. "
                       f"It can be read but never written — open it in Word and use "
                       f"Save As with the file type set to 'Word Document (.docx)'.")
    if fmt not in ("pdf", "docx"):
        return False, f"{path.name} is not a readable PDF or DOCX (detected: {fmt})."
    if fmt == "pdf" and path.suffix.lower() != ".pdf":
        return False, f"{path.name} is a PDF with a {path.suffix} extension."
    if fmt == "docx" and path.suffix.lower() != ".docx":
        return False, f"{path.name} is a DOCX with a {path.suffix} extension."
    return True, ""


def _pdf_head(path: Path, max_chars: int) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    total = 0
    # Page by page, stopping as soon as there is enough. Reading all 145 pages
    # of a real bid pack to look at the first 1500 characters takes seconds and
    # buys nothing.
    for page in reader.pages:
        text = page.extract_text() or ""
        if not text.strip():
            continue
        parts.append(text)
        total += len(text)
        if total >= max_chars:
            break
    return "\n".join(parts)[:max_chars]


def _docx_head(path: Path, max_chars: int) -> str:
    import docx

    document = docx.Document(str(path))
    parts: list[str] = []
    total = 0
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        parts.append(text)
        total += len(text)
        if total >= max_chars:
            return "\n".join(parts)[:max_chars]

    # Many SBD/MBD forms carry their heading inside the first table rather than
    # in a paragraph, so a paragraph-only reader sees an empty document.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if not text:
                    continue
                parts.append(text)
                total += len(text)
                if total >= max_chars:
                    return "\n".join(parts)[:max_chars]
    return "\n".join(parts)[:max_chars]


@dataclass
class HeadText:
    """The readable head of a document, and how it was obtained.

    `extract_text_head` returns only the string, which is all most callers
    want. This carries the extra fact the narration needs: whether the text
    came off the page or out of an OCR engine, because a user should be told
    when what they are reading is a machine's guess at a photograph.
    """

    text: str = ""
    ocr_used: bool = False
    ocr: "object | None" = None      # ocr.OcrResult when ocr_used, else None

    def __bool__(self) -> bool:
        return bool(self.text.strip())


def read_head(path: str | Path, max_chars: int = HEAD_CHARS) -> HeadText:
    """
    First `max_chars` of readable text, falling back to OCR on a scan.

    The fallback is only ever reached when the ordinary reader found nothing,
    so a document with a text layer is never sent to Vision — that would cost
    money to produce a worse copy of text already in the file.
    """
    path = Path(path)
    fmt = detect_format(path)
    try:
        if fmt == "pdf":
            text = _pdf_head(path, max_chars)
            if text.strip():
                return HeadText(text=text)
            return _ocr_head(path, max_chars)
        if fmt == "docx":
            return HeadText(text=_docx_head(path, max_chars))
    except Exception as exc:  # noqa: BLE001 - a malformed file must not crash a folder scan
        log.warning("agent_autofill.classification: could not read %s: %s", path.name, exc)
        return HeadText()
    return HeadText()


def _ocr_head(path: Path, max_chars: int) -> HeadText:
    """
    Last resort for a PDF that yielded no text: read it as an image.

    A failure here is reported, not swallowed. `HeadText.ocr` carries the
    reason so the caller can tell the user "this is a scan and OCR is not
    configured" instead of the older and much less useful "this document
    appears to be empty".
    """
    from agent_autofill.extraction import ocr as ocr_module

    if not ocr_module.needs_ocr(path):
        return HeadText()

    result = ocr_module.ocr_pdf(path)
    return HeadText(text=(result.text or "")[:max_chars],
                    ocr_used=True, ocr=result)


def extract_text_head(path: str | Path, max_chars: int = HEAD_CHARS) -> str:
    """First `max_chars` of readable text. Empty string if there is none."""
    return read_head(path, max_chars).text


_JSON_OBJECT = re.compile(r"\{.*\}", re.DOTALL)


def _parse_verdict(raw: str) -> dict | None:
    """
    Pull the JSON object out of the model's reply.

    Returns None rather than raising: an unparseable reply is treated as "do
    not proceed", which is the safe direction.
    """
    if not raw:
        return None
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\s*", "", text)
        text = re.sub(r"```\s*$", "", text).strip()
    match = _JSON_OBJECT.search(text)
    if not match:
        return None
    try:
        parsed = json.loads(match.group(0))
    except (ValueError, TypeError):
        return None
    return parsed if isinstance(parsed, dict) else None


def _coerce_confidence(value) -> float:
    try:
        conf = float(value)
    except (TypeError, ValueError):
        return 0.0
    return max(0.0, min(1.0, conf))


def classify_document(path: str | Path, company_id: str) -> ClassificationResult:
    """
    Decide whether `path` is worth running the autofill pipeline over.

    Does NOT check the tier or the daily quota — the orchestrator does that
    first, so that a company over its limit never reaches this function.
    """
    path = Path(path)
    result = ClassificationResult(path=str(path))

    ok, why = is_classifiable(path)
    if not ok:
        result.status = "unsupported_type"
        result.reason = why
        log.info("agent_autofill.classification skip (%s): %s", result.status, why)
        return result

    head_result = read_head(path)
    head = head_result.text
    result.text_head = head
    result.ocr_used = head_result.ocr_used
    if head_result.ocr_used and head_result.ocr is not None:
        from agent_autofill.extraction.ocr import ocr_note

        result.ocr_note = ocr_note(head_result.ocr, path.name)

    if len(head.strip()) < MIN_TEXT_CHARS:
        result.status = "unreadable"
        # Three different situations used to share one sentence: a scan nobody
        # could read, a scan OCR *could* have read if it were configured, and a
        # genuinely blank file. Telling them apart is the difference between
        # "this cannot be done" and "this needs one setting changed".
        result.reason = result.ocr_note or (
            f"{path.name} has no extractable text in its opening pages — "
            f"it is probably a scan. Agent Autofill cannot read it."
        )
        log.info("agent_autofill.classification skip (unreadable, ocr_used=%s): %s",
                 head_result.ocr_used, path.name)
        return result

    # Shared global throttle. Layer 3 in the existing scheme; there is
    # deliberately no second limiter for autofill.
    if not rate_limiter.check_global_rate_limit():
        result.status = "rate_limited"
        result.reason = ("The assistant is handling too many requests right now. "
                         "This document will be retried.")
        log.warning("agent_autofill.classification deferred (global rate limit): %s", path.name)
        return result

    user_prompt = (
        f"File name: {path.name}\n"
        f"--- opening text ---\n{head}\n--- end ---\n\n"
        "Classify it."
    )

    try:
        response = claude_client.call_claude_with_tracking(
            company_id=company_id,
            messages=[{"role": "user", "content": user_prompt}],
            system=SYSTEM_PROMPT,
            max_tokens=MAX_TOKENS,
            model=CLASSIFIER_MODEL,
        )
    except Exception as exc:  # noqa: BLE001 - one bad file must not stop a folder scan
        result.status = "error"
        result.api_called = True
        result.reason = f"Classification failed: {exc}"
        log.error("agent_autofill.classification error on %s: %s", path.name, exc)
        return result

    result.api_called = True
    verdict = _parse_verdict(response.get("content", ""))
    if verdict is None:
        result.status = "error"
        result.reason = "The classifier did not return a usable verdict; skipping this file."
        log.error("agent_autofill.classification unparseable reply for %s: %r",
                  path.name, (response.get("content") or "")[:200])
        return result

    result.is_tender = bool(verdict.get("is_tender"))
    result.confidence = _coerce_confidence(verdict.get("confidence"))
    result.reason = str(verdict.get("reason") or "").strip()[:300]
    result.document_type = str(verdict.get("document_type") or "").strip()[:80]

    log.info(
        "agent_autofill.classification %s: is_tender=%s confidence=%.2f threshold=%.2f "
        "proceed=%s type=%r reason=%r",
        path.name, result.is_tender, result.confidence, CONFIDENCE_THRESHOLD,
        result.proceed, result.document_type, result.reason,
    )
    return result
