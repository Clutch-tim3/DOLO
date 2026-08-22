"""Find fillable blanks in UNSTRUCTURED documents using layout geometry.

This is the common case for SA tender packs: a flat print with no AcroForm,
where "the field" is just visual space next to a printed label.

Everything here works off **coordinates**, never raw ``extract_text()`` output.
Raw text collapses the very information that identifies a blank — a run of
underscores and the eight-centimetre gap after "NAME OF BIDDER" both flatten
into ordinary whitespace, and column order across a two-column form scrambles.
``extract_words()`` / ``page.chars`` / ``find_tables()`` keep the x/y data that
makes the association rule possible.

Coordinate convention
---------------------
All bboxes are ``(x0, top, x1, bottom)`` in pdfplumber's **top-down** page
space: ``top`` grows downward from the top edge. This is deliberately *not*
PDF's native bottom-up space (which ``acroform_extractor`` reports, since that
is what the PDF stores). The fill engine must not mix the two.


THE ASSOCIATION RULE
====================
A blank's label is resolved by trying, in strict order:

  RULE 1 — SAME-LINE LEFT (preferred, highest confidence)
      The label is the run of words ending immediately to the left of the blank
      on the same text line. Two words are on the same line when their vertical
      spans overlap by at least ``LINE_OVERLAP_RATIO``. Walking left stops at:
        * a horizontal gap wider than ``LABEL_WORD_GAP`` (that is a different
          column, not more label), or
        * another blank on the same line (a second field's territory).
      The nearest word must start within ``MAX_LEFT_GAP`` of the blank, else the
      blank is treated as having no left label.

  RULE 2 — DIRECTLY ABOVE (fallback)
      Used when Rule 1 finds nothing — the "stacked" layout where a caption sits
      over its writing space. The candidate line must sit within
      ``MAX_ABOVE_GAP`` vertically and its horizontal span must overlap the
      blank's span by at least ``MIN_ABOVE_OVERLAP``. Column overlap is required
      precisely because a two-column form otherwise donates the left column's
      caption to the right column's blank.

  RULE 3 — NONE
      Reported with ``label_text=None`` and low confidence. It is never mapped
      to a canonical field. An unlabelled blank is a real finding, not a
      failure to hide.

Left beats above because SA forms are overwhelmingly "LABEL: ______" and
label-cell/value-cell tables. Above is the minority stacked-caption case, so
promoting it would mis-assign the common layout.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable, Sequence

import pdfplumber

# --------------------------------------------------------------------------
# Tuning constants — all in PDF points (1/72"). Derived from measurements on
# real SA tender documents, not guessed.
# --------------------------------------------------------------------------
MIN_UNDERSCORE_RUN = 3      # "___" is the shortest run that is reliably a blank
MIN_DOT_RUN = 5             # "...." is punctuation; 5+ is a leader
LEADER_CHAR_GAP = 2.5       # max x-gap between consecutive leader chars

#: A blank drawn as a rectangle rather than typed. See `_find_rule_runs`.
MAX_RULE_THICKNESS = 1.8    # above this it is a filled bar, not a rule
MIN_RULE_WIDTH = 28.0       # narrower than this holds nothing worth writing
RULE_TEXT_HEIGHT = 9.5      # the writing band sits above the rule
MAX_REPEATED_RULES = 3      # same x-range this often down a page = a column
LINE_TOLERANCE = 2.0        # vertical slack when clustering chars into a line
LINE_OVERLAP_RATIO = 0.45   # min vertical overlap to call two words same-line

LABEL_WORD_GAP = 14.0       # gap above which words stop being one label
MAX_LEFT_GAP = 65.0         # blank must start within this of its left label
MAX_ABOVE_GAP = 26.0        # vertical reach for the "label above" rule
MIN_ABOVE_OVERLAP = 0.30    # horizontal overlap ratio for the above rule

# Minimum writable width. At the 10-11pt type these forms use, average glyph
# advance is ~5.5pt, so 40pt is roughly 7 characters — below that a cell cannot
# hold a company name, registration number or date, and is grid structure rather
# than a field.
#
# Calibrated against the signature block on p85/p86 of BID_DOCUMENT_06FY27,
# where the "Date:" label column (x 339-374, 35pt wide) is empty on the rows
# that have no date. Those 4 spacer cells were being detected as fields, and two
# of them inherited "Capacity" from the far-left label cell and mapped to
# `capacity` at 100.0 — actionable phantom fields. Raising the floor from 22 to
# 40 removes all 4 and costs zero true positives on that document (the narrowest
# genuine field there is the 60pt dialling-code box on SBD 1).
MIN_BLANK_WIDTH = 40.0
MIN_TRAILING_GAP = 90.0     # trailing whitespace must be this wide to count
MAX_LABEL_CHARS = 120       # longer "labels" are prose, not field captions

_LEADER_CHARS = frozenset({"_", ".", "…", "·"})
_UNDERSCORE_CHARS = frozenset({"_"})
_INT_RE = re.compile(r"^\(?\d{1,4}\)?$")
_SENTENCE_TAIL_RE = re.compile(r"[.;]$")


@dataclass(frozen=True)
class Blank:
    """One detected fillable blank and the label associated with it.

    ``confidence`` is about *this blank being a real, writable field with this
    label* — it is independent of whether the label maps to a canonical field,
    which is ``field_alias_dictionary``'s job.
    """

    source: str                                   # "pdf" | "docx"
    page_number: int                              # 0-based
    bbox: tuple[float, float, float, float] | None  # (x0, top, x1, bottom)
    label_text: str | None
    label_bbox: tuple[float, float, float, float] | None
    confidence: float                             # 0.0 - 1.0
    strategy: str                                 # how it was detected
    label_origin: str                             # left | above | cell_left | cell_above | none
    notes: tuple[str, ...] = ()
    table_index: int | None = None
    row_index: int | None = None
    col_index: int | None = None

    @property
    def width(self) -> float:
        return 0.0 if self.bbox is None else self.bbox[2] - self.bbox[0]

    @property
    def position_str(self) -> str:
        if self.bbox is not None:
            x0, top, x1, bottom = self.bbox
            return f"({x0:.0f},{top:.0f})-({x1:.0f},{bottom:.0f})"
        if self.row_index is not None:
            return f"tbl{self.table_index} r{self.row_index} c{self.col_index}"
        return "-"


# --------------------------------------------------------------------------
# Geometry helpers
# --------------------------------------------------------------------------
def _vertical_overlap_ratio(a_top: float, a_bottom: float, b_top: float, b_bottom: float) -> float:
    overlap = min(a_bottom, b_bottom) - max(a_top, b_top)
    if overlap <= 0:
        return 0.0
    shorter = min(a_bottom - a_top, b_bottom - b_top)
    return overlap / shorter if shorter > 0 else 0.0


def _horizontal_overlap_ratio(a_x0: float, a_x1: float, b_x0: float, b_x1: float) -> float:
    overlap = min(a_x1, b_x1) - max(a_x0, b_x0)
    if overlap <= 0:
        return 0.0
    shorter = min(a_x1 - a_x0, b_x1 - b_x0)
    return overlap / shorter if shorter > 0 else 0.0


def _clean_label(text: str | None) -> str | None:
    """Trim a scraped label to its caption, or reject it as prose."""
    if text is None:
        return None
    cleaned = " ".join(text.replace("\n", " ").split())
    # The comma is stripped from BOTH ends. A blank in the middle of a sentence
    # picks up the comma that closed the previous clause: MBD 3.1's second blank
    # arrived as ", in my capacity as", which matches no alias because of one
    # leading character.
    cleaned = cleaned.strip(" \t:;,-–—_.")
    cleaned = " ".join(cleaned.split())
    if not cleaned:
        return None
    if len(cleaned) > MAX_LABEL_CHARS:
        return None
    # A caption is not a sentence. Require at least one alphanumeric char.
    if not any(ch.isalnum() for ch in cleaned):
        return None
    return cleaned


# --------------------------------------------------------------------------
# PDF: leader runs (underscores / dot leaders)
# --------------------------------------------------------------------------
def _find_leader_runs(page: pdfplumber.page.Page) -> list[dict[str, Any]]:
    """Group consecutive leader characters into runs, using char coordinates."""
    leaders = [
        c
        for c in page.chars
        if c.get("text") in _LEADER_CHARS
    ]
    if not leaders:
        return []

    leaders.sort(key=lambda c: (round(c["top"], 1), c["x0"]))

    runs: list[dict[str, Any]] = []
    current: list[dict[str, Any]] = []

    def flush() -> None:
        if not current:
            return
        chars = [c["text"] for c in current]
        is_underscore = all(ch in _UNDERSCORE_CHARS for ch in chars)
        minimum = MIN_UNDERSCORE_RUN if is_underscore else MIN_DOT_RUN
        if len(current) < minimum:
            return
        x0 = min(c["x0"] for c in current)
        x1 = max(c["x1"] for c in current)
        top = min(c["top"] for c in current)
        bottom = max(c["bottom"] for c in current)
        if x1 - x0 < MIN_BLANK_WIDTH:
            return
        runs.append(
            {
                "bbox": (x0, top, x1, bottom),
                "kind": "underscore" if is_underscore else "dot",
                "length": len(current),
            }
        )

    for char in leaders:
        if not current:
            current = [char]
            continue
        prev = current[-1]
        same_line = abs(char["top"] - prev["top"]) <= LINE_TOLERANCE
        adjacent = (char["x0"] - prev["x1"]) <= LEADER_CHAR_GAP
        # Mixing dots and underscores in one run is a formatting accident;
        # splitting keeps the run-length thresholds meaningful.
        same_kind = (char["text"] in _UNDERSCORE_CHARS) == (prev["text"] in _UNDERSCORE_CHARS)
        if same_line and adjacent and same_kind:
            current.append(char)
        else:
            flush()
            current = [char]
    flush()
    return runs


# --------------------------------------------------------------------------
# PDF: text lines
# --------------------------------------------------------------------------
def _build_lines(words: Sequence[dict[str, Any]]) -> list[dict[str, Any]]:
    """Cluster words into visual text lines."""
    if not words:
        return []
    ordered = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))
    lines: list[dict[str, Any]] = []
    current: list[dict[str, Any]] = [ordered[0]]

    for word in ordered[1:]:
        ref = current[-1]
        if _vertical_overlap_ratio(
            word["top"], word["bottom"], ref["top"], ref["bottom"]
        ) >= LINE_OVERLAP_RATIO:
            current.append(word)
        else:
            lines.append(_make_line(current))
            current = [word]
    lines.append(_make_line(current))
    return lines


def _make_line(words: list[dict[str, Any]]) -> dict[str, Any]:
    ws = sorted(words, key=lambda w: w["x0"])
    return {
        "words": ws,
        "top": min(w["top"] for w in ws),
        "bottom": max(w["bottom"] for w in ws),
        "x0": min(w["x0"] for w in ws),
        "x1": max(w["x1"] for w in ws),
    }


def _is_leader_word(word: dict[str, Any]) -> bool:
    text = word.get("text", "")
    return bool(text) and all(ch in _LEADER_CHARS for ch in text)


def _split_words_at_leaders(
    words: Sequence[dict[str, Any]],
    runs: Sequence[dict[str, Any]],
    page: pdfplumber.page.Page,
) -> list[dict[str, Any]]:
    """
    Cut a word where a leader run starts inside it.

    THE BUG THIS FIXES, on the owner's Johannesburg Water RFQ, MBD 3.1:

        Name of Bidder……………………………     Bid

    pdfplumber returns that as ONE word, "Bidder………………………", spanning
    x=116.8 to x=317.5 — the noun fused to fourteen U+2026 glyphs. The blank
    starts at x=149.5, so `_label_left` discards the word for overrunning the
    blank it labels, and the field arrives at the fill engine labelled

        "Name of"

    which matches nothing, so a form asking for the bidder's name came back
    blank while the company name sat in the profile. The owner: "the system
    leaves out too much info it shouldnt leave out."

    ONLY runs `_find_leader_runs` already accepted are cut at. That matters:
    "." is a leader character, so cutting at any leader char would split "No.",
    "2.1.1" and "Ltd." apart. The run thresholds — five dots, a contiguity gap —
    are what separate a leader from punctuation, and they are applied once,
    there, rather than reimplemented here.
    """
    boxes = [r["bbox"] for r in runs]
    if not boxes:
        return list(words)

    chars = page.chars
    out: list[dict[str, Any]] = []
    for word in words:
        text = word.get("text", "")
        starts = [
            b[0] for b in boxes
            if b[0] > word["x0"] + 0.5 and b[0] < word["x1"] - 0.5
            and _vertical_overlap_ratio(
                word["top"], word["bottom"], b[1], b[3]) > 0.3
        ]
        if not text or not starts:
            out.append(word)
            continue

        cut = min(starts)
        kept = [
            c for c in chars
            if c["x0"] >= word["x0"] - 0.6 and c["x1"] <= cut + 0.6
            and _vertical_overlap_ratio(
                word["top"], word["bottom"], c["top"], c["bottom"]) > 0.5
        ]
        kept = [c for c in kept if c["text"] not in _LEADER_CHARS]
        if not kept:
            # The whole word was leader. It is the blank, not a label.
            continue

        kept.sort(key=lambda c: c["x0"])
        trimmed = dict(word)
        trimmed["text"] = "".join(c["text"] for c in kept).strip()
        trimmed["x1"] = max(c["x1"] for c in kept)
        if trimmed["text"]:
            out.append(trimmed)

    return out


def _find_rule_runs(page: pdfplumber.page.Page) -> list[dict[str, Any]]:
    """
    Blanks drawn as a thin rectangle rather than typed as underscores.

    THE SECOND HALF of the same missing MBD 3.1 line:

        I (full name) ______________, in my capacity as ________, the duly
        authorized representative of ____________(company name)

    Not one of those blanks was extracted. They look like underscores and are
    not: `page.chars` holds a space and then nothing between x=111.8 and
    x=299.0. The rule is a filled rectangle 0.72pt high. `_find_leader_runs`
    reads characters, so it cannot see them, and three fields CairoAI holds
    values for — signatory name, signatory capacity, company name — were left
    blank on a page the owner then had to complete by hand.

    WHY THIS IS NARROW ON PURPOSE

    Every table border on the page is also a thin rectangle. Four filters keep
    those out, and a fifth lives at the call site:

      * thickness — a rule, not a filled bar;
      * width — wide enough to write in, and under 85% of the text width, since
        a rule spanning the page is a divider or a table edge;
      * not already inside a detected table cell (checked by the caller, which
        has the cell boxes);
      * A LABEL TO THE LEFT ON THE SAME LINE is required by the caller. This is
        the filter that does the work: a table border has prose above or below
        it, not a caption beside it.

    The returned box is the writing space ABOVE the rule, not the rule itself —
    a 0.72pt-high box has nowhere to put a value and no text line to pair with.
    """
    text_x0, text_x1 = page.bbox[0], page.bbox[2]
    max_width = (text_x1 - text_x0) * 0.85

    candidates: list[tuple[float, float, float]] = []
    for rect in page.rects:
        height = float(rect.get("height") or 0.0)
        if height > MAX_RULE_THICKNESS:
            continue
        candidates.append((float(rect["x0"]), float(rect["x1"]),
                           float(rect["top"])))
    for line in page.lines:
        if abs(float(line["top"]) - float(line["bottom"])) > MAX_RULE_THICKNESS:
            continue
        candidates.append((float(line["x0"]), float(line["x1"]),
                           float(line["top"])))

    kept = [(x0, x1, top) for x0, x1, top in candidates
            if MIN_RULE_WIDTH <= (x1 - x0) <= max_width]

    # A rule repeated at the same x-range down the page is a TABLE COLUMN, not
    # a field. Measured on the owner's Johannesburg Water RFQ: without this,
    # 30 of the 37 new blanks were rate cells in the Bill of Quantities —
    # "Disposal Per Ton", "(6m3) as specified" — stacked five and six deep at
    # identical widths. They are priced from the quotation and refused anyway,
    # so all they add is noise on the review screen, which is the failure the
    # extraction filter exists to prevent.
    columns: dict[tuple[int, int], int] = {}
    for x0, x1, _top in kept:
        columns[(round(x0), round(x1))] = columns.get((round(x0), round(x1)), 0) + 1

    runs: list[dict[str, Any]] = []
    for x0, x1, top in kept:
        if columns.get((round(x0), round(x1)), 0) >= MAX_REPEATED_RULES:
            continue
        runs.append({
            "kind": "rule",
            # The line sits under where a person writes, so the field is the
            # band above it.
            "bbox": (x0, top - RULE_TEXT_HEIGHT, x1, top + MAX_RULE_THICKNESS),
        })
    return runs


# --------------------------------------------------------------------------
# PDF: the association rule
# --------------------------------------------------------------------------
def _label_left(
    blank_bbox: tuple[float, float, float, float],
    lines: Sequence[dict[str, Any]],
    other_blanks: Sequence[tuple[float, float, float, float]],
) -> tuple[str | None, tuple[float, float, float, float] | None, float]:
    """RULE 1. Returns (label, label_bbox, gap)."""
    bx0, btop, bx1, bbottom = blank_bbox

    best_line = None
    best_overlap = 0.0
    for line in lines:
        overlap = _vertical_overlap_ratio(btop, bbottom, line["top"], line["bottom"])
        if overlap > best_overlap:
            best_overlap, best_line = overlap, line
    if best_line is None or best_overlap < LINE_OVERLAP_RATIO:
        return None, None, float("inf")

    # Words wholly to the left of the blank, nearest first.
    left_words = [
        w
        for w in best_line["words"]
        if w["x1"] <= bx0 + 1.0 and not _is_leader_word(w)
    ]
    if not left_words:
        return None, None, float("inf")
    left_words.sort(key=lambda w: w["x0"])

    gap = bx0 - left_words[-1]["x1"]
    if gap > MAX_LEFT_GAP:
        return None, None, gap

    # Walk left while words stay contiguous and we do not cross another blank.
    chosen: list[dict[str, Any]] = [left_words[-1]]
    for word in reversed(left_words[:-1]):
        prev = chosen[-1]
        if prev["x0"] - word["x1"] > LABEL_WORD_GAP:
            break
        crosses = any(
            ob is not blank_bbox
            and word["x1"] <= ob[0]
            and prev["x0"] >= ob[2]
            and _vertical_overlap_ratio(btop, bbottom, ob[1], ob[3]) >= LINE_OVERLAP_RATIO
            for ob in other_blanks
        )
        if crosses:
            break
        chosen.append(word)

    chosen.sort(key=lambda w: w["x0"])
    text = _clean_label(" ".join(w["text"] for w in chosen))
    if text is None:
        return None, None, gap
    lbox = (
        min(w["x0"] for w in chosen),
        min(w["top"] for w in chosen),
        max(w["x1"] for w in chosen),
        max(w["bottom"] for w in chosen),
    )
    return text, lbox, gap


def _label_above(
    blank_bbox: tuple[float, float, float, float],
    lines: Sequence[dict[str, Any]],
) -> tuple[str | None, tuple[float, float, float, float] | None, float]:
    """RULE 2. Returns (label, label_bbox, vertical_gap)."""
    bx0, btop, bx1, _ = blank_bbox
    best = None
    best_gap = float("inf")
    for line in lines:
        gap = btop - line["bottom"]
        if gap < 0 or gap > MAX_ABOVE_GAP:
            continue
        if _horizontal_overlap_ratio(bx0, bx1, line["x0"], line["x1"]) < MIN_ABOVE_OVERLAP:
            continue
        if gap < best_gap:
            best_gap, best = gap, line
    if best is None:
        return None, None, float("inf")

    words = [w for w in best["words"] if not _is_leader_word(w)]
    if not words:
        return None, None, best_gap
    text = _clean_label(" ".join(w["text"] for w in words))
    if text is None:
        return None, None, best_gap
    lbox = (best["x0"], best["top"], best["x1"], best["bottom"])
    return text, lbox, best_gap


#: Words a label ending in them is not a label — it is the middle of a sentence
#: that runs THROUGH the blank. The real caption is on the other side.
_CONNECTIVES = frozenset({
    "of", "as", "at", "by", "for", "to", "in", "on", "the", "a", "an", "and",
    "with", "from",
})


def _label_right_hint(
    blank_bbox: tuple[float, float, float, float],
    lines: Sequence[dict[str, Any]],
) -> tuple[str | None, tuple[float, float, float, float] | None]:
    """
    The parenthetical immediately after a blank, when the words before it are
    the middle of a sentence.

    THE CASE, from the owner's Johannesburg Water RFQ, MBD 3.1:

        ...the duly authorized representative of ______________(company name)

    Everything to the left of that blank is "representative of", which is not a
    caption — it is prose broken by the space you write in. The form's author
    knew that, which is why they put the answer's name on the RIGHT, in
    brackets, where a person reading the page will find it.

    So the label becomes "company name", which matches, and the field fills from
    the profile. Left as "representative of" it was discarded as a prose
    fragment and the owner completed his own company's name by hand.

    ONLY when the left-hand words end in a connective, and ONLY a parenthetical.
    "Signature of" plus "(Bidder)" would resolve to "Bidder" — still refused, by
    `never_fill_fields`, on a label the caller checks separately. Nothing here
    decides whether a field may be filled.
    """
    bx0, btop, bx1, bbottom = blank_bbox

    for line in lines:
        if _vertical_overlap_ratio(
                btop, bbottom, line["top"], line["bottom"]) < LINE_OVERLAP_RATIO:
            continue
        right = sorted(
            (w for w in line["words"]
             if w["x0"] >= bx1 - 1.0 and not _is_leader_word(w)),
            key=lambda w: w["x0"],
        )
        if not right:
            continue

        text = " ".join(w["text"] for w in right).strip()
        if not text.startswith("("):
            continue
        close = text.find(")")
        if close <= 1:
            continue

        inner = _clean_label(text[1:close])
        if inner is None:
            return None, None

        used = []
        for word in right:
            used.append(word)
            if ")" in word["text"]:
                break
        return inner, (
            min(w["x0"] for w in used), min(w["top"] for w in used),
            max(w["x1"] for w in used), max(w["bottom"] for w in used),
        )

    return None, None


def _ends_in_connective(label: str | None) -> bool:
    words = (label or "").strip().rstrip(":").split()
    return bool(words) and words[-1].lower() in _CONNECTIVES


def _score(
    strategy: str,
    origin: str,
    gap: float,
    label: str | None,
    width: float,
    notes: Sequence[str],
) -> float:
    """Confidence that this is a genuine writable field with this label."""
    if label is None:
        return 0.25

    base = {
        "underscore_run": 0.90,
        "dot_leader": 0.86,
        # A drawn rule is the same evidence as a typed underscore run, minus a
        # little: a table border can look like one, which is why the caller
        # requires a caption to its left before accepting it at all.
        "drawn_rule": 0.84,
        "ruled_cell": 0.88,
        "trailing_gap": 0.62,
        "docx_table_cell": 0.88,
        "docx_underscore_run": 0.85,
    }.get(strategy, 0.6)

    if origin in ("above", "cell_above"):
        base -= 0.12  # weaker evidence than a same-line label

    # Distance penalty: the further the label, the less certain the pairing.
    if gap != float("inf"):
        if gap > 40:
            base -= 0.10
        elif gap > 20:
            base -= 0.05

    # A trailing colon is a strong positive signal of a caption.
    if label.rstrip().endswith(":"):
        base += 0.04

    if len(label) <= 2:
        base -= 0.20
    if width and width < MIN_BLANK_WIDTH * 1.5:
        base -= 0.05

    for note in notes:
        if note in ("possible_toc_leader", "label_is_prose"):
            base -= 0.35

    return max(0.05, min(0.99, round(base, 3)))


# --------------------------------------------------------------------------
# PDF: ruled table cells
# --------------------------------------------------------------------------
def _extract_ruled_cells(page: pdfplumber.page.Page, page_number: int) -> list[Blank]:
    """Empty cells in a ruled table, labelled by the nearest filled cell.

    This is the dominant real pattern in SA SBD/MBD forms — the same
    label-cell/value-cell shape as the DOCX path.

    ``extract()`` distinguishes two kinds of falsy cell, and the difference
    matters: ``None`` means *no cell exists here* (a merged span's continuation
    columns), while ``''`` means *a real, empty cell*. Only ``''`` is a blank.
    Treating ``None`` as blank would invent a field per merged column.
    """
    blanks: list[Blank] = []
    try:
        tables = page.find_tables()
    except Exception:
        return blanks

    for t_index, table in enumerate(tables):
        try:
            grid = table.extract()
        except Exception:
            continue

        for r_index, row in enumerate(table.rows):
            if r_index >= len(grid):
                break
            text_row = grid[r_index]

            for c_index, cell_bbox in enumerate(row.cells):
                if cell_bbox is None or c_index >= len(text_row):
                    continue
                cell_text = text_row[c_index]
                # None => cell absent (merged continuation). Not a blank.
                if cell_text is None:
                    continue
                if cell_text.strip():
                    continue

                x0, top, x1, bottom = cell_bbox
                if (x1 - x0) < MIN_BLANK_WIDTH:
                    continue

                label = None
                label_bbox = None
                origin = "none"
                gap = float("inf")

                # Nearest non-empty cell to the LEFT in the same row.
                for back in range(c_index - 1, -1, -1):
                    if back >= len(text_row):
                        continue
                    candidate = text_row[back]
                    if candidate is None or not candidate.strip():
                        continue
                    cleaned = _clean_label(candidate)
                    if cleaned is None:
                        break
                    label = cleaned
                    origin = "cell_left"
                    cb = row.cells[back]
                    if cb is not None:
                        label_bbox = (cb[0], cb[1], cb[2], cb[3])
                        gap = x0 - cb[2]
                    break

                # Same guard as the DOCX path: a cell spanning essentially the
                # whole table width is a note/spacer band, so the row above is
                # not its label.
                table_width = table.bbox[2] - table.bbox[0]
                spans_full_row = table_width > 0 and (x1 - x0) >= 0.9 * table_width

                # Else the nearest non-empty cell ABOVE in the same column.
                if label is None and not spans_full_row:
                    for up in range(r_index - 1, -1, -1):
                        if up >= len(grid) or c_index >= len(grid[up]):
                            continue
                        candidate = grid[up][c_index]
                        if candidate is None or not candidate.strip():
                            continue
                        cleaned = _clean_label(candidate)
                        if cleaned is None:
                            break
                        label = cleaned
                        origin = "cell_above"
                        ub = table.rows[up].cells[c_index]
                        if ub is not None:
                            label_bbox = (ub[0], ub[1], ub[2], ub[3])
                            gap = top - ub[3]
                        break

                notes: list[str] = []
                if label is None:
                    notes.append("no_label_in_row_or_column")
                if spans_full_row:
                    notes.append("full_row_merge")

                blanks.append(
                    Blank(
                        source="pdf",
                        page_number=page_number,
                        bbox=(x0, top, x1, bottom),
                        label_text=label,
                        label_bbox=label_bbox,
                        confidence=_score("ruled_cell", origin, gap, label, x1 - x0, notes),
                        strategy="ruled_cell",
                        label_origin=origin,
                        notes=tuple(notes),
                        table_index=t_index,
                        row_index=r_index,
                        col_index=c_index,
                    )
                )
    return blanks


# --------------------------------------------------------------------------
# PDF: trailing whitespace before a line break
# --------------------------------------------------------------------------
def _extract_trailing_gaps(
    lines: Sequence[dict[str, Any]],
    page_number: int,
    right_margin: float,
    covered: Sequence[tuple[float, float, float, float]],
) -> list[Blank]:
    """A caption ending in ':' followed by empty space to the right margin.

    OFF BY DEFAULT — measured precision on real data is 0%.

    Across pages [6, 46, 47, 48, 56, 57, 61, 84, 85, 86] of BID_DOCUMENT_06FY27
    this strategy produced 10 detections and 0 real fields. Every one was one of:

      * a wrapped label line — p6 "No:" is the tail of "CENTRAL SUPPLIER
        DATABASE No:", whose actual blank is the underscore run on the line
        above and was already detected;
      * a section heading — p86 "Notes:";
      * a question whose answer space is the dotted lines *below* it, already
        detected as a dot leader — p47 "2.2.1 If so, furnish particulars:";
      * ordinary prose that happens to end in a colon.

    The trailing colon is too weak a signal: SA tender prose uses colons
    constantly. Keeping the code (the spec asks for the strategy, and a form
    built from tab stops rather than leader glyphs would need it) but requiring
    callers to opt in via ``include_trailing_gaps=True``.
    """
    blanks: list[Blank] = []
    for line in lines:
        words = [w for w in line["words"] if not _is_leader_word(w)]
        if not words:
            continue
        last = max(words, key=lambda w: w["x1"])
        raw_text = " ".join(w["text"] for w in sorted(words, key=lambda w: w["x0"]))

        if not raw_text.rstrip().endswith(":"):
            continue

        gap_width = right_margin - last["x1"]
        if gap_width < MIN_TRAILING_GAP:
            continue

        bbox = (last["x1"] + 2.0, line["top"], right_margin, line["bottom"])
        # Do not double-report space already claimed by a leader run or cell.
        if any(_horizontal_overlap_ratio(bbox[0], bbox[2], c[0], c[2]) > 0.5
               and _vertical_overlap_ratio(bbox[1], bbox[3], c[1], c[3]) > 0.5
               for c in covered):
            continue

        label = _clean_label(raw_text)
        if label is None:
            continue

        notes: list[str] = []
        if len(raw_text) > 80:
            notes.append("label_is_prose")

        blanks.append(
            Blank(
                source="pdf",
                page_number=page_number,
                bbox=bbox,
                label_text=label,
                label_bbox=(line["x0"], line["top"], last["x1"], line["bottom"]),
                confidence=_score("trailing_gap", "left", 2.0, label, gap_width, notes),
                strategy="trailing_gap",
                label_origin="left",
                notes=tuple(notes),
            )
        )
    return blanks


# --------------------------------------------------------------------------
# PDF entry point
# --------------------------------------------------------------------------
def extract_pdf_blanks(
    path: str | Path,
    pages: Iterable[int] | None = None,
    include_trailing_gaps: bool = False,
) -> list[Blank]:
    """Detect blanks across a PDF using word/char geometry.

    ``pages`` is an optional iterable of 0-based page indices.
    ``include_trailing_gaps`` defaults to False — see
    ``_extract_trailing_gaps`` for the measured precision that motivated it.
    """
    results: list[Blank] = []
    wanted = set(pages) if pages is not None else None

    with pdfplumber.open(str(path)) as pdf:
        for page_number, page in enumerate(pdf.pages):
            if wanted is not None and page_number not in wanted:
                continue

            try:
                words = page.extract_words(keep_blank_chars=False, use_text_flow=False)
            except Exception:
                continue
            # Leader runs are found BEFORE the lines are built, because a word
            # fused to a leader ("Bidder…………") has to be cut at the run before
            # it can serve as a label. See `_split_words_at_leaders`.
            leader_runs = _find_leader_runs(page)
            words = _split_words_at_leaders(words, leader_runs, page)
            lines = _build_lines(words)

            cell_blanks = _extract_ruled_cells(page, page_number)
            cell_boxes = [b.bbox for b in cell_blanks if b.bbox]

            runs = leader_runs + _find_rule_runs(page)
            run_boxes = [r["bbox"] for r in runs]

            leader_blanks: list[Blank] = []
            for run in runs:
                bbox = run["bbox"]
                # A leader inside a detected table cell is that cell's blank;
                # reporting both would double-count the same field.
                if any(
                    _horizontal_overlap_ratio(bbox[0], bbox[2], cb[0], cb[2]) > 0.6
                    and _vertical_overlap_ratio(bbox[1], bbox[3], cb[1], cb[3]) > 0.5
                    for cb in cell_boxes
                ):
                    continue

                label, lbox, gap = _label_left(bbox, lines, run_boxes)
                origin = "left"

                # "...representative of ______(company name)". The words to the
                # left are the middle of a sentence; the form names the answer
                # on the right, in brackets. Prefer that.
                from_parenthetical = False
                if label is not None and _ends_in_connective(label):
                    hint, hbox = _label_right_hint(bbox, lines)
                    if hint:
                        label, lbox, gap = hint, hbox, 2.0
                        from_parenthetical = True

                if label is None:
                    if run["kind"] == "rule":
                        # A drawn rule with nothing beside it is almost always a
                        # table border or a divider, not a field. Requiring a
                        # same-line caption is what makes this strategy safe to
                        # turn on at all — see `_find_rule_runs`.
                        continue
                    label, lbox, gap = _label_above(bbox, lines)
                    origin = "above" if label else "none"

                notes: list[str] = []
                if from_parenthetical:
                    notes.append("label_from_parenthetical")
                if run["kind"] == "dot" and _looks_like_toc(bbox, lines):
                    notes.append("possible_toc_leader")

                strategy = {
                    "underscore": "underscore_run",
                    "rule": "drawn_rule",
                }.get(run["kind"], "dot_leader")
                leader_blanks.append(
                    Blank(
                        source="pdf",
                        page_number=page_number,
                        bbox=bbox,
                        label_text=label,
                        label_bbox=lbox,
                        confidence=_score(
                            strategy, origin, gap, label, bbox[2] - bbox[0], notes
                        ),
                        strategy=strategy,
                        label_origin=origin,
                        notes=tuple(notes),
                    )
                )

            results.extend(cell_blanks)
            results.extend(leader_blanks)

            if include_trailing_gaps:
                right_margin = max((w["x1"] for w in words), default=0.0)
                covered = cell_boxes + run_boxes
                results.extend(
                    _extract_trailing_gaps(lines, page_number, right_margin, covered)
                )

            # THE PAGE IS DONE. RELEASE IT.
            #
            # pdfplumber caches every derived object on the page — chars, words,
            # edges, the table finder's output — and holds them for the life of
            # the PDF object. Nothing here needs a page after its blanks are
            # collected, so on a long document the cache is pure growth.
            #
            # Measured on the owner's 145-page pack: extracting ONE document
            # peaked at 573 MB and settled at 195 MB. The Cloud Run function has
            # 1024 MiB, and a pack holds several documents, so submitting one
            # killed the container and the browser saw a 503. Every 503 in the
            # logs is /api/autofill-packs/<id>/submit, each paired with a
            # "Memory limit of 1024 MiB exceeded" seconds later.
            #
            # `flush_cache` is pdfplumber's own remedy for exactly this.
            try:
                page.flush_cache()
            except Exception:  # noqa: BLE001 - freeing memory must never fail a fill
                pass

    results.sort(key=lambda b: (b.page_number, b.bbox[1] if b.bbox else 0, b.bbox[0] if b.bbox else 0))
    # P0-4: drop what is not a field before anything downstream sees it.
    # Every one of these was already being refused, so no fill changes — what
    # goes is 534 [ ! ] markers on the page and 534 lines in the flag list.
    return [b for b in results if is_fillable_candidate(b)]


def _looks_like_toc(
    bbox: tuple[float, float, float, float], lines: Sequence[dict[str, Any]]
) -> bool:
    """A dot leader followed by a bare page number is a contents entry.

    Table-of-contents rows look exactly like fill-in blanks to a geometric
    detector. This is the single biggest false-positive source for dot leaders.
    """
    x1, top, bottom = bbox[2], bbox[1], bbox[3]
    for line in lines:
        if _vertical_overlap_ratio(top, bottom, line["top"], line["bottom"]) < LINE_OVERLAP_RATIO:
            continue
        after = [w for w in line["words"] if w["x0"] >= x1 - 1.0 and not _is_leader_word(w)]
        if len(after) == 1 and _INT_RE.match(after[0]["text"].strip()):
            return True
    return False


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------
def extract_docx_blanks(path: str | Path) -> list[Blank]:
    """Detect blanks in a .docx: empty table cells and underscore runs.

    The SA MBD/SBD form pattern is a two-column table — a label cell beside an
    empty cell to be completed. Same association rule as the PDF path (left,
    then above), applied to the table grid.

    Structural notes that changed the implementation:

    * python-docx returns the *same* ``_tc`` element repeatedly for a merged
      cell, once per grid position it spans. Without de-duplication by element
      identity, one merged cell is reported as several blanks.
    * There are no coordinates in a .docx — layout is resolved at render time.
      ``bbox`` is therefore ``None`` and position is reported as
      (table, row, column). Anything downstream that assumes a bbox must handle
      this; it is the main structural difference between the two paths.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(str(path))
    blanks: list[Blank] = []

    for t_index, table in enumerate(document.tables):
        grid: list[list[tuple[str, Any] | None]] = []
        for row in table.rows:
            row_cells: list[tuple[str, Any] | None] = []
            try:
                cells = row.cells
            except (IndexError, ValueError):
                grid.append(row_cells)
                continue
            for cell in cells:
                row_cells.append((cell.text or "", cell._tc))
            grid.append(row_cells)

        seen_elements: set[int] = set()

        for r_index, row_cells in enumerate(grid):
            for c_index, entry in enumerate(row_cells):
                if entry is None:
                    continue
                text, element = entry
                if text.strip():
                    continue
                # Merged cell repeated across grid positions -> report once.
                key = id(element)
                if key in seen_elements:
                    continue
                seen_elements.add(key)

                label = None
                origin = "none"
                for back in range(c_index - 1, -1, -1):
                    prev = row_cells[back]
                    if prev is None or not prev[0].strip():
                        continue
                    cleaned = _clean_label(prev[0])
                    if cleaned is None:
                        break
                    label, origin = cleaned, "cell_left"
                    break

                # A cell merged across the whole row is a note/spacer band, not
                # the value half of a label/value pair. Inheriting the label
                # from the row above turns it into a phantom field — on the MBD
                # fixture the trailing merged row inherited "BBBEE Status Level
                # of Contribution" and mapped at 100.0, which would have written
                # the B-BBEE level into a free-text note area.
                spans_full_row = (
                    len(row_cells) > 1
                    and len({id(e) for entry in row_cells if entry for _, e in (entry,)}) == 1
                )

                if label is None and not spans_full_row:
                    for up in range(r_index - 1, -1, -1):
                        if c_index >= len(grid[up]):
                            continue
                        above = grid[up][c_index]
                        if above is None or not above[0].strip():
                            continue
                        cleaned = _clean_label(above[0])
                        if cleaned is None:
                            break
                        label, origin = cleaned, "cell_above"
                        break

                notes: list[str] = []
                if label is None:
                    notes.append("no_label_in_row_or_column")
                if spans_full_row:
                    notes.append("full_row_merge")

                blanks.append(
                    Blank(
                        source="docx",
                        page_number=0,  # .docx has no page geometry pre-render
                        bbox=None,
                        label_text=label,
                        label_bbox=None,
                        confidence=_score(
                            "docx_table_cell", origin, 0.0, label, 0.0, notes
                        ),
                        strategy="docx_table_cell",
                        label_origin=origin,
                        notes=tuple(notes),
                        table_index=t_index,
                        row_index=r_index,
                        col_index=c_index,
                    )
                )

    # Underscore runs in body paragraphs ("Name: ______").
    for p_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text or ""
        for match in re.finditer(r"_{%d,}" % MIN_UNDERSCORE_RUN, text):
            before = text[: match.start()]
            label = _clean_label(before.split("  ")[-1] if before.strip() else None)
            notes: list[str] = []
            if label is None:
                notes.append("no_label_before_run")
            blanks.append(
                Blank(
                    source="docx",
                    page_number=0,
                    bbox=None,
                    label_text=label,
                    label_bbox=None,
                    confidence=_score(
                        "docx_underscore_run",
                        "left" if label else "none",
                        0.0,
                        label,
                        0.0,
                        notes,
                    ),
                    strategy="docx_underscore_run",
                    label_origin="left" if label else "none",
                    notes=tuple(notes),
                    table_index=None,
                    row_index=p_index,
                    col_index=match.start(),
                )
            )

    return blanks


# --- rejecting things that are not fields -------------------------------------
#
# P0-4. On the owner's 145-page pack the extractor proposed 651 blanks. Eight
# of them fill. 534 are refused as unreadable or unrecognised — and every one
# of those draws a red [ ! ] on the page, so the document came back with 534
# markers on it and a flag list nobody could work through.
#
# They are not fields. They are the points-allocation table an organ of state
# completes before issuing the tender (80, 20, 100), fragments of the
# surrounding instructions ("of this tender", "the tenderer)"), and blanks with
# no readable label at all.
#
# NOTHING HERE CHANGES WHAT GETS FILLED. Every blank rejected below was already
# being skipped — measured, not assumed: the same 8 fields fill before and
# after. What goes is the noise: the markers on the page and the flags in the
# list.
#
# CONSERVATIVE ON PURPOSE. "Name of Bidder (tenderer)" fills successfully and
# ends in a bracket; "Value of work inclusive of VAT (Rand)" is a real column
# header. So a rule as crude as "ends with )" would drop real fields, and the
# rules below are written to reject only what is unambiguously prose or
# pre-printed data.

#: Words that are a whole label only when something has been mis-split. A blank
#: labelled "of" is a fragment of a sentence, not a question.
_FUNCTION_WORDS = frozenset({
    "of", "and", "or", "the", "a", "an", "to", "in", "on", "at", "by", "for",
    "with", "as", "is", "are", "be", "been", "that", "this", "these", "those",
    "i", "we", "it", "if", "no", "yes", "not", "from", "per",
})

#: Openings that only occur mid-declaration. These are the sworn-statement
#: preambles on SBD 4 and SBD 6.1, not field labels.
_PROSE_OPENINGS = (
    "i, the undersigned", "do hereby", "i hereby", "we hereby",
    # "in my capacity" was here. It is prose, and it is also the only thing
    # printed beside the blank where the signatory's capacity goes — the SBD
    # signature block runs its question through the blank rather than in front
    # of it. It now maps to `capacity` in the alias dictionary, so rejecting it
    # here dropped a field the profile can fill.
    "representative of", "on behalf of",
    "it is a condition", "failure to", "note:", "please note",
    "kindly note", "the tenderer", "each tenderer", "bidders are",
    "tenderers are", "in terms of", "subject to", "provided that",
)

#: An instruction that has been cut off. A real label does not end mid-clause.
_FRAGMENT_ENDINGS = (",", ";", ".]", ".)", " -", "–", "—", ":-")


def _is_prose_fragment(label: str) -> bool:
    """Whether a label is a piece of the form's prose rather than a question."""
    text = (label or "").strip()
    if not text:
        return False

    lowered = text.lower()

    # A bare function word is always a mis-split.
    if lowered.strip(".:,)") in _FUNCTION_WORDS:
        return True

    # Declaration preambles.
    if lowered.startswith(_PROSE_OPENINGS):
        return True

    # Cut off mid-clause. Checked before the length rule because some of these
    # are long.
    if text.endswith(_FRAGMENT_ENDINGS):
        return True

    # Starts mid-sentence: a lower-case opening that is not a known field
    # phrasing. Real labels on these forms are either capitalised or title
    # case; "of this tender" is not.
    if text[0].islower() and " " in text:
        return True

    return False


def _is_preprinted_value(label: str) -> bool:
    """
    Whether a label is the form's own pre-printed data rather than a prompt.

    On SBD 6.1 the points-allocation table carries 80, 20 and 100 — the
    preference split, filled in by the organ of state before the tender is
    issued. Offering those as blanks invites a supplier to overwrite the
    evaluation criteria.
    """
    text = (label or "").strip()
    if not text:
        return False

    stripped = text.replace(" ", "").replace(".", "").replace(",", "")
    stripped = stripped.replace("%", "").replace("R", "").replace("-", "")
    if stripped.isdigit():
        return True

    return False


def is_fillable_candidate(blank) -> bool:
    """
    Whether this detected blank should be offered as a field at all.

    A blank that is not a candidate is dropped before it reaches the fill
    engine, so it never becomes a [ ! ] on the page or a line in the flag list.
    """
    label = (getattr(blank, "label_text", "") or "").strip()

    # No label, nowhere to go. 164 of the owner's 651 were these, and the fill
    # engine could only ever refuse them with "could not read what this field
    # is for" — which is true and useless.
    if not label:
        return False

    # A label the form's author put in brackets beside the blank — "(company
    # name)", "(full name)" — is not prose that happens to be lower case. It is
    # the author naming the answer, which is stronger evidence than any rule
    # below. Without this exemption the prose test rejected "company name" on
    # MBD 3.1 and the owner wrote his own company's name in by hand.
    if "label_from_parenthetical" in (getattr(blank, "notes", None) or ()):
        return True

    if _is_prose_fragment(label) and not _is_known_field_label(label):
        return False

    if _is_preprinted_value(label):
        return False

    return True


def _is_known_field_label(label: str) -> bool:
    """
    Whether the alias dictionary recognises this label as a named field.

    The prose test rejects any label that starts lower-case and has a space in
    it — a good rule, because "of this tender" and "the tenderer)" are sentence
    fragments the extractor picks up off busy pages.

    It is also wrong about the SBD signature block, which runs its question
    THROUGH the blank rather than in front of it:

        I (full name) ______, in my capacity as ______, the duly authorized
        representative of ______(company name)

    "in my capacity as" starts lower-case and has spaces, so it was discarded
    as prose — and it is the only thing printed beside the blank where the
    signatory's capacity goes. The owner filled his own job title in by hand.

    The dictionary is the authority on what is a field label. If it names one,
    a heuristic about capitalisation should not overrule it. This cannot widen
    what gets FILLED beyond what the dictionary already maps, and every label
    that survives still goes through `is_blocked` and `decide` in the fill
    engine — a recognised label is not permission to write anything.

    AN EXACT ALIAS ONLY, never a fuzzy one. The first version accepted any
    match and a test caught it immediately: "representative of (tenderer)" is a
    declaration preamble, and the fuzzy matcher scores it 95 against
    contact_person on the strength of the word "representative". So prose that
    resembles a field would have been readmitted as a field — the exact failure
    the prose test exists to prevent. An exact hit means the label IS an entry
    in the dictionary, which somebody wrote down on purpose.
    """
    try:
        from agent_autofill.extraction import match_label

        match = match_label(label)
    except Exception:  # noqa: BLE001 - a heuristic must never break extraction
        return False

    return (getattr(match, "status", None) == "exact"
            and bool(getattr(match, "canonical", None)))
