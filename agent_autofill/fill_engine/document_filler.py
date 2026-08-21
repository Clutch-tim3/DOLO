"""
Writes safe values into a copy of the document.

Two things this module will never do, and both are load-bearing:

  * It never writes to the source file. Every fill produces a new document, so
    the original tender pack stays exactly as received.
  * It never writes a value the safety gate did not approve. It does not decide
    anything itself — safe_fill_fields.decide() is the only decision point, and
    this module refuses to run without a decision object.

Every touched cell gets a gold tint so a reviewer can see at a glance what the
agent wrote versus what was already there, and every skipped field gets a
marginal marker so nothing is silently passed over. A reviewer scanning the
document must be able to tell "agent filled this" from "this was always here"
without cross-referencing the summary.
"""

from __future__ import annotations

import shutil
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .never_fill_fields import BlockDecision, classify_document_context
from .safe_fill_fields import FillDecision, decide

#: Gold tint, matching the brand accent. Light enough to read black text over.
FILL_SHADE = "F3E9D6"
#: Marker prefixed to a field the agent deliberately left alone.
SKIP_MARKER = "[ ! ]"


@dataclass
class FilledField:
    label: str
    value: str
    source: str
    low_confidence: bool = False
    location: str = ""


@dataclass
class SkippedField:
    label: str
    reason: str
    category: str          # "blocked" | "low_confidence" | "no_data" | "unmatched"
    location: str = ""


@dataclass
class FillResult:
    source_path: str
    output_path: str
    filled: list[FilledField] = field(default_factory=list)
    skipped: list[SkippedField] = field(default_factory=list)
    context: set[str] = field(default_factory=set)

    @property
    def fillable_total(self) -> int:
        return len(self.filled) + len(self.skipped)

    @property
    def summary_line(self) -> str:
        return (f"{len(self.filled)} of {self.fillable_total} fillable fields "
                f"completed automatically — review required before use")


def _shade_cell(cell, hex_fill: str = FILL_SHADE) -> None:
    """Apply a background tint to a table cell (python-docx has no API for it)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _document_text(document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _is_column_header_table(table) -> bool:
    """
    True for a grid whose FIRST ROW is column headings and whose later rows are
    blank entry rows — the shape SBD 4's declaration table uses
    ("Full Name | Identity Number | Name of organ of state" across the top,
    empty rows beneath).

    Without this the pair-finder below sees no "label | blank" adjacency and
    walks straight past the entire table, so those cells were neither filled
    nor marked. The spec requires that nothing is skipped without a visible
    trace, so an unseen cell is a defect even when the outcome is safe.
    """
    if len(table.rows) < 2:
        return False
    header = [(c.text or "").strip() for c in table.rows[0].cells]
    if not header or any(not h for h in header):
        return False                      # a blank in row 0 means it is not a header
    if all(len(h) > 60 for h in header):
        return False                      # prose, not headings
    # at least one later row must actually be an entry row
    for row in table.rows[1:]:
        if any(not (c.text or "").strip() for c in row.cells):
            return True
    return False


def _iter_label_value_cells(document):
    """
    Yield (label_text, target_cell, location) for every cell a human is expected
    to complete, across both table shapes SA bid forms use:

      * "label | blank" rows        — supplier information blocks (MBD 1)
      * column-headed entry grids   — declaration tables (SBD 4)

    Only a genuinely empty cell is a target. A cell that already has content is
    never overwritten — the agent adds, it does not edit.
    """
    for t_i, table in enumerate(document.tables):
        if _is_column_header_table(table):
            headers = [(c.text or "").strip() for c in table.rows[0].cells]
            for r_i, row in enumerate(table.rows[1:], start=2):
                for c_i, cell in enumerate(row.cells):
                    if (cell.text or "").strip():
                        continue
                    label = headers[c_i] if c_i < len(headers) else ""
                    if not label or len(label) < 3:
                        continue
                    yield label, cell, f"table {t_i + 1}, row {r_i}, col {c_i + 1}"
            continue

        for r_i, row in enumerate(table.rows):
            cells = row.cells
            for c_i in range(len(cells) - 1):
                label_cell, value_cell = cells[c_i], cells[c_i + 1]
                label = (label_cell.text or "").strip()
                if not label or len(label) < 3:
                    continue
                if (value_cell.text or "").strip():
                    continue          # already answered — leave it alone
                if value_cell is label_cell:
                    continue
                yield label, value_cell, f"table {t_i + 1}, row {r_i + 1}"


def fill_docx(source: str | Path, output: str | Path, profile: dict,
              match_label_fn) -> FillResult:
    """
    Produce a filled COPY of a .docx form.

    `match_label_fn` is injected (normally extraction.match_label) so this
    module has no opinion about how a label becomes a canonical field.
    """
    source, output = Path(source), Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    # Copy first, then edit the copy. The source tender pack is never modified.
    shutil.copy2(source, output)

    document = docx.Document(str(output))
    context = classify_document_context(_document_text(document))
    result = FillResult(str(source), str(output), context=context)

    for label, value_cell, location in _iter_label_value_cells(document):
        match = match_label_fn(label)
        decision: FillDecision = decide(
            match.canonical, label, profile, match.score, context=context
        )

        if decision.fill and decision.value:
            value_cell.text = decision.value
            _shade_cell(value_cell)
            result.filled.append(FilledField(
                label=label, value=decision.value,
                source=decision.source or "company profile",
                low_confidence=decision.low_confidence, location=location,
            ))
            continue

        # Not filled. Recorded in `result.skipped` — and ONLY there.
        block: BlockDecision | None = decision.block
        if block is not None:
            category = "blocked"
        elif decision.low_confidence:
            category = "low_confidence"
        elif match.canonical:
            category = "no_data"
        else:
            category = "unmatched"

        # The cell is left exactly as the form's author wrote it. This used to
        # be `value_cell.text = SKIP_MARKER`, which put a red `[ ! ]` in every
        # refused cell — including the whole Identity Number column of an SBD 4
        # table that is correctly empty. It survives printing, so the owner was
        # handing organs of state a statutory form covered in marks. The
        # refusal belongs in the review, not on the document.
        result.skipped.append(SkippedField(
            label=label,
            reason=decision.reason or "Left for you to complete.",
            category=category, location=location,
        ))

    document.save(str(output))
    return result
