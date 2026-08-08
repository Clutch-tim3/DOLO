"""
Review state, written into the file itself.

The problem this solves: a draft that has been half-reviewed and a draft that
has been fully reviewed are byte-for-byte similar and look identical when
someone forwards one by email. The database knows which is which; the recipient
does not, and the recipient is the one who submits it.

So the state goes in the document. Both ways, deliberately:

  1. **Core properties** (`content_status`, `category`, `comments`, `keywords`,
     `subject`, `last_modified_by`). These are machine-readable, survive a copy
     or an email attachment, and Word surfaces `content_status` as "Status" and
     `category`/`comments` in File > Info. This is what a program — ours or
     anyone's — can check without parsing the body.

  2. **A visible banner paragraph at the top of the body.** Because nobody opens
     File > Info. A person who receives the document, scrolls, and prints it
     must see the words "UNREVIEWED DRAFT" without being told to look for them.

Neither is sufficient alone: properties are invisible to humans, and a banner is
trivially deleted and invisible to code. Writing both means the two have to be
defeated separately, and a mismatch between them is itself a signal — which is
why `read_review_state()` returns both and flags disagreement.

Nothing written here ever says "approved", "final", or "submission-ready". The
strongest thing this module will stamp on a document is "REVIEWED DRAFT", and
even that carries the sentence that the marked fields still need completing by
hand.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

import docx
from docx.shared import Pt, RGBColor

#: Every line the banner writes begins with this, so a re-stamp can find and
#: remove the previous banner instead of stacking a second one on top.
BANNER_SENTINEL = "CAIROAI AGENT AUTOFILL"

#: Machine-readable prefix inside `keywords`. Grep-able, stable, versioned.
KEYWORD_PREFIX = "CAIROAI_AUTOFILL_V1"

STATUS_UNREVIEWED = "UNREVIEWED DRAFT"
STATUS_REVIEWED = "REVIEWED DRAFT"

_VALID_STATUSES = (STATUS_UNREVIEWED, STATUS_REVIEWED)

#: OPC caps every core property string at 255 characters and python-docx raises
#: ValueError rather than truncating. A long source path or company name would
#: otherwise abort a stamp, so every property written here is clipped.
MAX_PROP_CHARS = 255


def _clip(text: str, limit: int = MAX_PROP_CHARS) -> str:
    text = " ".join((text or "").split())
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


class ReviewStateError(ValueError):
    """Raised when a caller tries to describe a document as more finished than it is."""


@dataclass(frozen=True)
class ReviewStamp:
    """
    What gets written into the document.

    The invariant is enforced in `__post_init__`, not at the call site: a
    REVIEWED stamp with open flags cannot be constructed at all. That means
    there is no in-process object a caller can hand to `stamp_docx()` that
    would mark a document reviewed while something is still outstanding — they
    would have to lie about `flags_open`, which is fabrication rather than a
    gap in the gate.
    """

    review_id: str
    status: str
    flags_total: int
    flags_open: int
    filled_count: int = 0
    company_name: str = ""
    source_document: str = ""
    acknowledged: list[tuple[str, str, str]] = field(default_factory=list)
    stamped_at: str = ""

    def __post_init__(self):
        if self.status not in _VALID_STATUSES:
            raise ReviewStateError(
                f"Unknown review status {self.status!r}. "
                f"Only {_VALID_STATUSES} may be written into a document."
            )
        if self.flags_open < 0 or self.flags_total < 0:
            raise ReviewStateError("Flag counts cannot be negative.")
        if self.flags_open > self.flags_total:
            raise ReviewStateError(
                f"flags_open ({self.flags_open}) exceeds flags_total ({self.flags_total})."
            )
        if self.status == STATUS_REVIEWED and self.flags_open != 0:
            raise ReviewStateError(
                f"Refusing to build a '{STATUS_REVIEWED}' stamp with {self.flags_open} "
                "flagged field(s) still unacknowledged. Acknowledge each one first."
            )
        if not self.stamped_at:
            object.__setattr__(self, "stamped_at", datetime.now().isoformat(timespec="seconds"))

    # --- rendered forms ---------------------------------------------------

    @property
    def keyword_line(self) -> str:
        return (
            f"{KEYWORD_PREFIX} status={self.status.replace(' ', '_')} "
            f"flags_open={self.flags_open} flags_total={self.flags_total} "
            f"acknowledged={len(self.acknowledged)} filled={self.filled_count} "
            f"review_id={self.review_id} stamped_at={self.stamped_at}"
        )

    @property
    def headline(self) -> str:
        if self.status == STATUS_REVIEWED:
            return f"{BANNER_SENTINEL} — {STATUS_REVIEWED} — NOT YET A SUBMISSION"
        return f"{BANNER_SENTINEL} — {STATUS_UNREVIEWED} — DO NOT SUBMIT"

    @property
    def detail(self) -> str:
        if self.status == STATUS_REVIEWED:
            return (
                f"All {self.flags_total} flagged field(s) were acknowledged by a person "
                f"on {self.stamped_at}. {self.filled_count} field(s) were pre-filled by "
                "CairoAI from the company profile and are shaded. Every field marked "
                "[ ! ] must still be completed by hand. No signature has been applied "
                "and no pricing has been entered. Review reference "
                f"{self.review_id}."
            )
        return (
            f"{self.flags_open} of {self.flags_total} flagged field(s) have NOT been "
            "reviewed by a person. This document is an incomplete draft. "
            f"{self.filled_count} field(s) were pre-filled by CairoAI from the company "
            "profile and are shaded; every field marked [ ! ] is still blank. Do not "
            "submit this document and do not treat any value in it as verified. "
            f"Review reference {self.review_id}."
        )

    def properties_comment(self) -> str:
        """
        The `comments` core property: one dense sentence, inside 255 characters.

        The per-field acknowledgement log deliberately does NOT go here. It does
        not fit, and appending an audit appendix to a municipality's bid form
        would mean telling the user to delete part of the document before
        submitting — a footgun worse than the problem. The log lives in
        `autofill_review_item`; the document carries the state, which is what a
        recipient needs to know.
        """
        if self.status == STATUS_REVIEWED:
            body = (
                f"REVIEWED DRAFT. All {self.flags_total} flagged field(s) acknowledged "
                f"individually by a person on {self.stamped_at}. {self.filled_count} "
                "field(s) pre-filled by CairoAI. Fields marked [ ! ] must still be "
                "completed by hand. No signature applied, no pricing entered. "
                "NOT A SUBMISSION."
            )
        else:
            body = (
                f"UNREVIEWED DRAFT. {self.flags_open} of {self.flags_total} flagged "
                f"field(s) NOT reviewed by a person. {self.filled_count} field(s) "
                "pre-filled by CairoAI; fields marked [ ! ] are blank. "
                "DO NOT SUBMIT."
            )
        return _clip(f"{body} Review {self.review_id}.")


# --- writing ---------------------------------------------------------------


def _strip_previous_banner(document) -> int:
    """
    Remove any banner this module wrote before.

    Only the first few block elements are examined: a banner is always inserted
    at the very top, and scanning the whole body would let a sentinel string
    that happens to appear in the form's own text delete real content.
    """
    body = document.element.body
    removed = 0
    for para in list(document.paragraphs[:4]):
        if (para.text or "").strip().startswith(BANNER_SENTINEL):
            body.remove(para._p)
            removed += 1
    return removed


def _insert_banner(document, stamp: ReviewStamp) -> None:
    """
    Put the banner at the very top of the body.

    `add_paragraph` appends, so the new element is moved to index 0 afterwards.
    This works whether the document starts with a paragraph or a table, which
    `paragraphs[0].insert_paragraph_before()` does not — several SA bid forms
    open directly with a table and would raise IndexError there.
    """
    para = document.add_paragraph()

    head = para.add_run(stamp.headline)
    head.bold = True
    head.font.size = Pt(11)
    head.font.color.rgb = (
        RGBColor(0x8A, 0x6D, 0x3B) if stamp.status == STATUS_REVIEWED
        else RGBColor(0xC0, 0x1F, 0x1F)
    )

    para.add_run().add_break()

    body_run = para.add_run(stamp.detail)
    body_run.font.size = Pt(8.5)
    body_run.italic = True

    document.element.body.insert(0, para._p)


def stamp_docx(path: str | Path, stamp: ReviewStamp) -> dict:
    """
    Write `stamp` into the .docx at `path`, in place.

    Returns what was written, so a caller can log it without re-reading.
    """
    path = Path(path)
    if path.suffix.lower() != ".docx":
        raise ReviewStateError(
            f"Review state can only be stamped into a .docx; got {path.suffix!r}. "
            "Legacy .doc is read-only and PDF stamping is not built."
        )

    document = docx.Document(str(path))
    _strip_previous_banner(document)
    _insert_banner(document, stamp)

    props = document.core_properties
    # `content_status` is Word's own "Status" field — the closest thing the
    # format has to a purpose-built slot for this, and it shows in File > Info.
    props.content_status = _clip(stamp.status)
    props.category = _clip(f"CairoAI Agent Autofill — {stamp.status}")
    props.subject = _clip(stamp.headline)
    props.comments = stamp.properties_comment()
    props.keywords = _clip(stamp.keyword_line)
    props.last_modified_by = "CairoAI Agent Autofill (draft only)"

    document.save(str(path))
    return {
        "path": str(path),
        "status": stamp.status,
        "flags_open": stamp.flags_open,
        "flags_total": stamp.flags_total,
        "keywords": stamp.keyword_line,
    }


# --- reading back ----------------------------------------------------------

_KW_RE = re.compile(r"(\w+)=(\S+)")


def read_review_state(path: str | Path) -> dict:
    """
    Read the review state back out of a saved .docx.

    Returns both channels and whether they agree. `stamped` is False for any
    document this module never touched, which is the correct answer for a form
    that arrived from a municipality.
    """
    path = Path(path)
    document = docx.Document(str(path))
    props = document.core_properties

    keywords = props.keywords or ""
    parsed: dict[str, str] = {}
    if keywords.startswith(KEYWORD_PREFIX):
        parsed = dict(_KW_RE.findall(keywords))

    banner = ""
    for para in document.paragraphs[:4]:
        text = (para.text or "").strip()
        if text.startswith(BANNER_SENTINEL):
            banner = text
            break

    prop_status = props.content_status or ""
    banner_status = ""
    if banner:
        # "UNREVIEWED DRAFT" CONTAINS "REVIEWED DRAFT". Testing for the reviewed
        # status first therefore reports every unreviewed draft as reviewed —
        # which this function did, turning the one check meant to catch a forged
        # file into a source of false assurance. Longest match first, and the
        # word boundary is checked so no future status can be a substring of
        # another and reintroduce it.
        for candidate in sorted(_VALID_STATUSES, key=len, reverse=True):
            if re.search(rf"(?<![A-Z]){re.escape(candidate)}", banner):
                banner_status = candidate
                break

    stamped = bool(parsed) or bool(banner)

    def _int(key: str):
        try:
            return int(parsed[key])
        except (KeyError, ValueError):
            return None

    return {
        "stamped": stamped,
        "content_status": prop_status,
        "category": props.category or "",
        "subject": props.subject or "",
        "keywords": keywords,
        "comments": props.comments or "",
        "banner_present": bool(banner),
        "banner_text": banner,
        "banner_status": banner_status,
        "review_id": parsed.get("review_id"),
        "flags_open": _int("flags_open"),
        "flags_total": _int("flags_total"),
        "acknowledged": _int("acknowledged"),
        "filled": _int("filled"),
        "stamped_at": parsed.get("stamped_at"),
        # A document whose properties claim REVIEWED while the visible banner
        # says otherwise (or is missing) has been tampered with or partially
        # rewritten. Surfacing it beats silently trusting the friendlier half.
        "channels_agree": (not stamped) or (prop_status == banner_status),
    }
