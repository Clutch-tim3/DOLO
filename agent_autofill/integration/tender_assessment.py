"""
Eligibility and win-probability, on the document the autofill just read.

When a tender is classified, the user should not have to go to another tab and
re-upload the same PDF to find out whether they are even allowed to bid. This
module runs the existing pipeline against the same file and returns its result
in the same shape `/api/predict` already returns, so the Agent chat and the
frontend render it without a second code path.

Everything here is a call into code that already exists:

    predict.eligibility_gate.check_hard_eligibility   — the hard gate
    models.pdf_parser.parse_tender_document           — feature extraction
    predict.predict.*                                 — the ML pipeline
    models.sa_scoring.*                               — PPPFA scoring / uplift
    predict.regional_router.predict_tender_region     — ZA/UK routing
    app.inject_parsed_features                        — the route's own feature merge

`inject_parsed_features` lives only in app.py, so it is imported from there
rather than copied. The import is deferred to call time: app.py imports the
agent stack at module scope, and a top-level import here would close the loop.

One thing worth stating plainly, because it changes results: the eligibility
gate's verdict depends on the supplier profile it is given. `/api/predict`
passes province/municipality as "Unknown" and assumes CSD, CIDB and tax
clearance are held. This module keeps those exact defaults and only overrides a
field when the company profile actually carries a value for it, so a tender that
disqualifies through the route disqualifies here too. A company with tender
history or a matching locality legitimately gets a different answer — the gate
is about that company, not about the document alone.
"""

from __future__ import annotations

from pathlib import Path

from models.pdf_parser import extract_text_from_pdf, parse_tender_document
from models.sa_scoring import (
    adjust_probability_for_sa,
    calculate_total_sa_score,
    get_bbbee_recommendation,
    get_evaluation_system,
)
from predict.eligibility_gate import check_hard_eligibility
from predict.predict import (
    build_new_features,
    encode_and_impute,
    extract_features_from_tender_id,
    predict,
)
from predict.regional_router import detect_region, predict_tender_region

#: `/api/predict` imputes these when the bid document does not supply them.
#: Kept identical so the two surfaces cannot disagree about the same tender.
DEFAULT_SUPPLIER_PRICE = 450000.0
DEFAULT_COMPETITORS = 4
DEFAULT_BBBEE_LEVEL = 9


def _document_text(path: Path) -> str:
    """
    Text for the eligibility gate, whatever the document is.

    Format comes from `legacy_doc_reader.detect_format`, i.e. from magic bytes,
    never from the extension — seven of the sa_forms fixtures are OLE2 `.doc`
    behind a `.docx` name, and python-docx raises PackageNotFoundError on those.
    """
    from agent_autofill.extraction.legacy_doc_reader import detect_format, read_doc

    fmt = detect_format(path)

    if fmt == "pdf":
        return extract_text_from_pdf(path) or ""

    if fmt == "doc":
        return read_doc(path).text or ""

    if fmt == "docx":
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    return path.read_text(errors="ignore")


def build_supplier_profile(company_profile: dict | None, pit_total_wins=None) -> dict:
    """
    The dict `check_hard_eligibility` expects, built from a company profile.

    Defaults mirror `/api/predict` exactly. A registration is assumed present
    unless the profile says otherwise, because the alternative — inferring
    "no CSD" from an empty profile field — manufactures hard failures for
    companies that simply have not filled the wizard in.
    """
    profile = company_profile or {}
    return {
        "pit_total_wins": (
            pit_total_wins if pit_total_wins is not None
            else profile.get("pit_total_wins", 0)
        ),
        "province": profile.get("province") or "Unknown",
        "registered_municipality": (
            profile.get("registered_municipality")
            or profile.get("municipality")
            or "Unknown"
        ),
        "has_csd": profile.get("has_csd", True),
        "has_cidb": profile.get("has_cidb", True),
        "has_tax_clearance": profile.get("has_tax_clearance", True),
    }


def _bbbee_int(company_profile: dict | None, override=None) -> int:
    """
    B-BBEE level as the scoring code wants it.

    The stored column holds strings like 'Level 1 Contributor' — BUILD_STATE
    flags the type mismatch — so the digit is pulled out rather than cast.
    """
    if override is not None:
        try:
            return int(override)
        except (TypeError, ValueError):
            pass
    raw = (company_profile or {}).get("bbbee_level")
    if raw is None:
        return DEFAULT_BBBEE_LEVEL
    if isinstance(raw, int):
        return raw
    import re

    m = re.search(r"\d+", str(raw))
    return int(m.group()) if m else DEFAULT_BBBEE_LEVEL


def assess_tender(
    document_path: str | Path,
    company_profile: dict | None = None,
    artifacts: dict | None = None,
    model_version: str = "sailor",
    supplier_price: float | None = None,
    bbbee_level: int | None = None,
) -> dict:
    """
    Eligibility + win probability for one tender document.

    Returns the `/api/predict` response shape with an added `eligibility` block
    and a `region` block. When the ML half cannot run (a .docx form pack, a
    document the parser cannot read, missing artifacts) the eligibility half is
    still returned and `prediction_available` is False with a stated reason —
    a partial answer beats an exception, and beats a fabricated probability.
    """
    path = Path(document_path)
    if not path.exists():
        return {"status": "error", "message": f"Document not found: {path}"}

    text = _document_text(path)
    bbbee_to_use = _bbbee_int(company_profile, bbbee_level)
    supplier_name = (company_profile or {}).get("company_name") or "NEW COMPANY SA"

    parsed: dict = {}
    parse_error = None
    if path.suffix.lower() == ".pdf":
        try:
            parsed = parse_tender_document(path) or {}
        except Exception as e:  # a malformed PDF must not take the gate down
            parse_error = f"{type(e).__name__}: {e}"
    else:
        parse_error = (
            f"The win-probability model reads tender PDFs; {path.name} is "
            f"{path.suffix or 'an unknown format'}. Eligibility was still checked."
        )

    tender_value = parsed.get("tender_value")

    # --- eligibility (always) ------------------------------------------------
    supplier_profile = build_supplier_profile(company_profile)
    eligibility = check_hard_eligibility(text, supplier_profile)
    disqualified = not eligibility["eligible"]
    hard_failures = [f["reason"] for f in eligibility["hard_failures"]]
    logistics_warnings = [w["reason"] for w in eligibility["logistics_warnings"]]

    region = predict_tender_region(
        {"currency": "ZAR" if detect_region(text) == "ZA" else ""},
        text_context=text,
    )

    price = supplier_price if supplier_price is not None else parsed.get("bid_price")
    if price is None:
        price = DEFAULT_SUPPLIER_PRICE

    # None when the document does not state a competing price, which is the
    # normal case for a sealed-bid tender. It used to fall back to
    # `price * 0.88`, and `pdf_parser` handed up `bid_price * 0.9` before that,
    # so the PPPFA price score was almost always computed against a number
    # derived from the bidder's own price. `calculate_total_sa_score` withholds
    # the score when this is None.
    lowest_price = parsed.get("lowest_price")

    base = {
        "status": "success",
        "document": str(path),
        "supplier_name": supplier_name,
        "bbbee_level": bbbee_to_use,
        "disqualified": disqualified,
        "recommendation": "DISQUALIFIED" if disqualified else None,
        "hard_failures": hard_failures,
        "logistics_warnings": logistics_warnings,
        "eligibility": {
            "eligible": eligibility["eligible"],
            "hard_failures": eligibility["hard_failures"],
            "logistics_warnings": eligibility["logistics_warnings"],
            "max_achievable_functionality_pct": eligibility["max_achievable_functionality_pct"],
            "confidence": eligibility["confidence"],
        },
        "region": region,
        "parsed_tender_value": tender_value,
        "extraction_completeness": parsed.get("extraction_completeness"),
        "prediction_available": False,
        "prediction_unavailable_reason": None,
        "win_probability": None,
        "base_probability": None,
        "sa_adjusted_probability": None,
        "threshold": None,
        "confidence": None,
        "sa_analysis": None,
    }

    # A hard failure ends it. `/api/predict` returns 0.0 probabilities on the
    # disqualified path rather than a model score, because a win probability
    # for a bid that will be thrown out on compliance is a number that only
    # misleads. Same choice here.
    if disqualified:
        eval_sys = get_evaluation_system(tender_value)
        base.update({
            "win_probability": 0.0,
            "base_probability": 0.0,
            "sa_adjusted_probability": 0.0,
            "confidence": "PASS",
            "prediction_unavailable_reason": (
                "Disqualified by the hard eligibility gate; the model was not run."
            ),
            "sa_analysis": {
                "evaluation_system": eval_sys,
                "bbbee_advice": get_bbbee_recommendation(bbbee_to_use),
                "parsed_supplier_price": price,
                "parsed_lowest_price": lowest_price,
                "parsed_tender_value": tender_value,
                "base_probability": None,
                "final_probability": None,
            },
        })
        return base

    if parse_error or not parsed:
        base["prediction_unavailable_reason"] = parse_error or "Nothing parsed from the document."
        return base

    if artifacts is None:
        try:
            import app as _app  # deferred: app.py imports the agent stack itself

            artifacts = (
                _app.artifacts_conquest if model_version == "conquest"
                else _app.artifacts_sailor
            )
        except Exception as e:
            base["prediction_unavailable_reason"] = f"Model artifacts unavailable: {e}"
            return base

    if not artifacts:
        base["prediction_unavailable_reason"] = "Model artifacts are not loaded."
        return base

    try:
        from app import inject_parsed_features

        import re as _re
        import uuid as _uuid

        id_match = _re.search(r"\b(us_[0-9]{8})\b", text)
        tender_id = id_match.group(1) if id_match else str(_uuid.uuid4())

        # app.py builds this once at startup as get_feature_list(metadata);
        # same call, same artifacts dict.
        from predict.predict import get_feature_list

        feature_list = get_feature_list(artifacts["metadata"])

        features_df = extract_features_from_tender_id(
            tender_id, supplier_name, feature_list, artifacts["medians"]
        )
        features_df = build_new_features(features_df, artifacts["medians"])
        features_df = inject_parsed_features(features_df, parsed, price)
        features_df = encode_and_impute(
            features_df, artifacts["encoder"], artifacts["cat_cols"], artifacts["medians"]
        )

        if artifacts["xgb_model"].feature_names is not None:
            for feat in artifacts["xgb_model"].feature_names:
                if feat not in features_df.columns:
                    features_df[feat] = artifacts["medians"].get(feat, 0)
            features_df = features_df[artifacts["xgb_model"].feature_names]

        pred_res = predict(artifacts, features_df, mock_supplier_name=supplier_name)
        base_prob = pred_res["probability"]

        sa_score = calculate_total_sa_score(
            supplier_price=price,
            lowest_competing_price=lowest_price,
            bbbee_level=bbbee_to_use,
            tender_value_zar=tender_value,
            num_competitors=DEFAULT_COMPETITORS,
        )
        sa_adj = adjust_probability_for_sa(
            base_probability=base_prob,
            sa_score_dict=sa_score,
            num_competitors=DEFAULT_COMPETITORS,
        )
        final_probability = sa_adj["final_probability"]
        threshold = artifacts["threshold"]

        # No PPPFA score means no adjusted probability, and a recommendation
        # derived from a probability we do not have is the fabrication in a
        # different costume. The eligibility verdict, the evaluation system and
        # the B-BBEE points below are unaffected — they never depended on it.
        if final_probability is None:
            recommendation = None
            confidence = None
        else:
            recommendation = "PURSUE" if final_probability >= threshold else "PASS"

            if final_probability > threshold + 0.15:
                confidence = "HIGH"
            elif final_probability > threshold + 0.05:
                confidence = "MEDIUM"
            elif final_probability > threshold:
                confidence = "LOW"
            else:
                confidence = "PASS"

        base.update({
            "prediction_available": final_probability is not None,
            "prediction_unavailable_reason": sa_adj.get("unavailable_reason"),
            "tender_identifier": tender_id,
            "win_probability": final_probability,
            "base_probability": base_prob,
            "sa_adjusted_probability": final_probability,
            "recommendation": recommendation,
            "confidence": confidence,
            "threshold": threshold,
            "sa_analysis": {
                "evaluation_system": sa_score["evaluation_system"],
                "price_score": sa_score["price_score"],
                "price_score_available": sa_score["price_score_available"],
                "price_score_unavailable_reason": sa_score["price_score_unavailable_reason"],
                "bbbee_points": sa_score["bbbee_points"],
                "max_bbbee_points": sa_score["max_bbbee_points"],
                "total_score": sa_score["total_score"],
                "competitive_position": sa_score["competitive_position"],
                "base_probability": base_prob,
                "final_probability": final_probability,
                "adjusted_probability": final_probability,
                "uplift": sa_adj["uplift"],
                "bbbee_advice": get_bbbee_recommendation(bbbee_to_use),
                "parsed_supplier_price": price,
                "parsed_lowest_price": lowest_price,
                "parsed_tender_value": tender_value,
            },
        })
    except Exception as e:
        base["prediction_unavailable_reason"] = f"{type(e).__name__}: {e}"

    return base
