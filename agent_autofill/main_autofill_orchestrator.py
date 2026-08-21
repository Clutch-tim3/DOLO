"""
The pipeline, wired end to end — and stopped short of the finish line.

    trigger -> tier check -> classify -> extract -> fill -> review summary -> STOP

The last arrow is the point of this module. It ends at a human confirmation
gate and does not step over it. Nothing here marks a draft reviewed, approved
or exportable; `reviewed` and `exportable` are returned as False and there is
no code path that sets them True. That decision belongs to the export gate
(Subagent 6), which is a separate reviewer with separate rules.

Two things this module deliberately does NOT do:

  * It makes no fill decisions. Every field goes through
    fill_engine.safe_fill_fields.decide() via document_filler.fill_docx. This
    module never inspects a label, never resolves a value, and never writes
    into a document itself. If it did, there would be two decision points and
    only one of them would be tested.

  * It never widens what may be filled. No signature is applied, no price is
    written, no declaration is answered. Those refusals live in the fill
    engine and this module simply reports them.

Ordering is load-bearing. The tier and quota check runs FIRST — before the
file is opened and long before Haiku is called — because its whole purpose is
to avoid paying for a request that would be refused. The quota is consumed
LAST, and only when a draft was actually produced. The quotation flow learned
that the hard way: consuming on entry meant a failed run still burned the
user's allowance for the day.
"""

from __future__ import annotations

import logging
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path

from agent import file_paths
from agent.memory.company_store import get_company_profile
from agent.subscription import check_autofill_quota, get_company_tier, log_autofill_run
from agent_autofill.classification.is_tender_document import (
    CONFIDENCE_THRESHOLD,
    ClassificationResult,
    classify_document,
)
from agent_autofill.extraction import extract_document, match_label
from agent_autofill.extraction.legacy_doc_reader import (
    LegacyDocError,
    detect_format,
    iter_label_blank_pairs,
    read_doc,
)
from agent_autofill.fill_engine.document_filler import FillResult, fill_docx
from agent_autofill.fill_engine.review_summary import render_html

log = logging.getLogger("agent_autofill.orchestrator")

#: Terminal states. Only ``awaiting_review`` means a draft exists, and even
#: that one is not usable until a human has been through it.
STATUS_AWAITING_REVIEW = "awaiting_review"          # draft written, human must confirm
STATUS_ANALYSIS_ONLY = "analysis_only"              # tender PDF: reported, not writable
STATUS_REFUSED_TIER = "refused_tier"                # plan does not include autofill
STATUS_REFUSED_QUOTA = "refused_quota"              # daily limit reached
STATUS_LEGACY_DOC = "legacy_doc_read_only"          # OLE2 .doc — readable, never writable
STATUS_SKIPPED_UNSUPPORTED = "skipped_unsupported"  # not a PDF or DOCX
STATUS_SKIPPED_UNREADABLE = "skipped_unreadable"    # no text layer
STATUS_SKIPPED_NOT_TENDER = "skipped_not_tender"    # below the confidence threshold
STATUS_DEFERRED = "deferred_rate_limited"
STATUS_ERROR = "error"

#: Statuses that consume one of the company's daily autofills. A refusal, a
#: skip, or a failure must never appear here.
_CONSUMES_QUOTA = {STATUS_AWAITING_REVIEW, STATUS_ANALYSIS_ONLY}

_SAFE_NAME = re.compile(r"[^A-Za-z0-9._-]+")


@dataclass
class AutofillRun:
    """
    The whole outcome of one document, including the parts that did not happen.

    `reviewed` and `exportable` are here so that callers have somewhere honest
    to read them from, and they are always False. This module has no setter for
    either.
    """

    company_id: str
    source_path: str
    status: str
    message: str = ""
    tier: str = ""
    quota: dict = field(default_factory=dict)
    quota_consumed: bool = False
    claude_calls: int = 0

    classification: ClassificationResult | None = None
    extraction_summary: dict | None = None
    fill_result: FillResult | None = None

    output_document: str | None = None
    review_summary_path: str | None = None
    review_summary_url: str | None = None

    #: The gate. Always True / False / False out of this module.
    requires_human_confirmation: bool = True
    reviewed: bool = False
    exportable: bool = False

    @property
    def produced_draft(self) -> bool:
        return self.status == STATUS_AWAITING_REVIEW and self.output_document is not None

    def to_dict(self) -> dict:
        """A JSON-safe view for an API response or a log line."""
        cls = self.classification
        return {
            "company_id": self.company_id,
            "source": Path(self.source_path).name,
            "status": self.status,
            "message": self.message,
            "tier": self.tier,
            "quota": self.quota,
            "quota_consumed": self.quota_consumed,
            "claude_calls": self.claude_calls,
            "classification": None if cls is None else {
                "is_tender": cls.is_tender,
                "confidence": round(cls.confidence, 3),
                "threshold": CONFIDENCE_THRESHOLD,
                "document_type": cls.document_type,
                "reason": cls.reason,
                "status": cls.status,
            },
            "extraction": self.extraction_summary,
            "fill": None if self.fill_result is None else {
                "filled": len(self.fill_result.filled),
                "skipped": len(self.fill_result.skipped),
                "fillable_total": self.fill_result.fillable_total,
                "summary_line": self.fill_result.summary_line,
                "context": sorted(self.fill_result.context),
            },
            "output_document": self.output_document,
            "review_summary_url": self.review_summary_url,
            # Never negotiable, and never set anywhere in this module.
            "requires_human_confirmation": True,
            "reviewed": False,
            "exportable": False,
        }


def _output_name(source: Path, suffix: str) -> str:
    """
    A generated filename that survives file_paths.safe_generated_path().

    Real tender file names contain spaces, brackets and slashes
    ("REVISED SBD 4 -Annexure A (1).docx"), none of which that validator
    accepts, so the stem is slugged and a short unique tag added rather than
    reusing the original name.
    """
    stem = _SAFE_NAME.sub("_", source.stem).strip("._-") or "document"
    return f"autofill_{stem[:60]}_{uuid.uuid4().hex[:8]}{suffix}"


def _legacy_doc_report(path: Path) -> tuple[str, dict | None]:
    """
    Read-only handling for a binary .doc.

    There is no pure-Python writer for the format, so the honest outcome is a
    clear instruction plus whatever analysis is possible — never silence, and
    never an empty draft that looks like a failed fill.
    """
    instruction = (
        f"{path.name} is a legacy Word document (.doc). Agent Autofill can read it "
        f"but cannot write to it — there is no way to produce a filled copy of this "
        f"format. Open it in Word, choose File > Save As, and set the file type to "
        f"'Word Document (.docx)', then re-run. Renaming the file does not convert it."
    )
    try:
        doc = read_doc(path)
    except LegacyDocError as exc:
        return f"{instruction}\n\n(It could not be analysed either: {exc})", None

    pairs = list(iter_label_blank_pairs(doc))
    analysis = {
        "readable": True,
        "table_rows": len(doc.rows),
        "label_blank_pairs": len(pairs),
        "writable": doc.writable,  # always False
    }
    if pairs:
        instruction += (
            f"\n\nFor reference, {len(pairs)} label/blank pairs were found in it, so "
            f"the .docx version should be fillable."
        )
    return instruction, analysis


def _document_text(path, doc_type: str) -> str:
    """The document's words, for reading facts the tender states about itself."""
    try:
        if doc_type == "pdf":
            import fitz

            with fitz.open(str(path)) as doc:
                return "\n".join(page.get_text() for page in doc)
        if doc_type == "docx":
            import docx

            document = docx.Document(str(path))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.extend(c.text for c in row.cells)
            return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001 - reading is best effort
        log.warning("could not read text from %s: %s", path, exc)
    return ""


def _with_preference_claim(profile: dict, report) -> tuple[dict, dict]:
    """
    Add SBD 6.1's points claim to the profile for THIS document only.

    `SBD_COMPLIANCE.md` P0-2: Level 1 under 80/20 is a claim of 20 points, and
    "leaving it blank costs points for no reason".

    It is not stored. The claim is true of one tender, not of the company — the
    same Level 1 company claims 20 points on an 80/20 tender and 10 on a 90/10
    one — so it goes onto the dict handed to the fill engine and nowhere else.
    Writing it to the profile would carry one tender's answer onto the next.

    Both inputs must be known or nothing is added, and `preference` records why
    so the user can be asked. A blank claim costs points; a wrong one is a
    misrepresentation to an organ of state.
    """
    from agent_autofill.fill_engine.preference_points import (
        detect_preference_system,
        goal_claims,
        points_claim,
    )

    profile = dict(profile or {})
    text = _document_text(report.path, report.doc_type)
    system, evidence = detect_preference_system(text)
    claim = points_claim(profile.get("bbbee_level"), system)

    if claim is not None:
        # Whole numbers on the form: SBD 6.1's table is 20, 18, 14 — never 20.0.
        profile["bbbee_points_claim"] = (
            str(int(claim)) if float(claim).is_integer() else str(claim))

    return profile, {
        "system": system,
        "evidence": evidence,
        "bbbee_level": profile.get("bbbee_level"),
        "points_claim": claim,
        "goals": goal_claims(profile),
    }


def _pdf_analysis(report) -> dict:
    """Extraction counts for a tender PDF, which the fill engine cannot write."""
    counts = report.summary()
    counts["strategy_used"] = report.strategy_used
    counts["pages_with_blanks"] = len({
        f.blank.page_number for f in report.fields if f.blank.page_number is not None
    })
    return counts


def run_autofill(
    company_id: str,
    source_path: str | Path,
    *,
    output_dir: str | Path | None = None,
) -> AutofillRun:
    """
    Run one document through the pipeline and stop at the confirmation gate.

    Never raises for an ordinary bad document — a folder scan must survive a
    corrupt file. Programmer errors (a missing company_id) still raise.
    """
    if not company_id:
        raise ValueError("company_id is required — autofill is always tenant-scoped.")

    source = Path(source_path)
    tier = get_company_tier(company_id)

    # ---- 1. Tier + quota. First, always, before anything is opened. --------
    quota = check_autofill_quota(company_id)
    if not quota.get("allowed"):
        status = (STATUS_REFUSED_TIER if quota.get("code") == "tier_disabled"
                  else STATUS_REFUSED_QUOTA)
        log.info("agent_autofill.orchestrator refused (%s) company=%s file=%s",
                 status, company_id, source.name)
        return AutofillRun(
            company_id=company_id, source_path=str(source), status=status,
            message=quota.get("reason", "Agent Autofill is not available."),
            tier=tier, quota=quota, claude_calls=0,
        )

    if not source.exists():
        return AutofillRun(
            company_id=company_id, source_path=str(source), status=STATUS_ERROR,
            message=f"File not found: {source.name}", tier=tier, quota=quota,
        )

    # ---- 2. Legacy .doc, by magic bytes rather than extension. -------------
    # Checked before classification so an unwritable file never costs an API
    # call, and so the user gets the actionable instruction rather than a
    # generic "unsupported type".
    if detect_format(source) == "doc":
        message, analysis = _legacy_doc_report(source)
        log.info("agent_autofill.orchestrator legacy .doc: %s", source.name)
        return AutofillRun(
            company_id=company_id, source_path=str(source), status=STATUS_LEGACY_DOC,
            message=message, tier=tier, quota=quota, extraction_summary=analysis,
            claude_calls=0,
        )

    # ---- 3. Classification. The only paid step. ----------------------------
    classification = classify_document(source, company_id)
    calls = 1 if classification.api_called else 0

    if classification.status == "unsupported_type":
        return AutofillRun(
            company_id=company_id, source_path=str(source),
            status=STATUS_SKIPPED_UNSUPPORTED, message=classification.reason,
            tier=tier, quota=quota, classification=classification, claude_calls=calls,
        )
    if classification.status == "unreadable":
        return AutofillRun(
            company_id=company_id, source_path=str(source),
            status=STATUS_SKIPPED_UNREADABLE, message=classification.reason,
            tier=tier, quota=quota, classification=classification, claude_calls=calls,
        )
    if classification.status == "rate_limited":
        return AutofillRun(
            company_id=company_id, source_path=str(source), status=STATUS_DEFERRED,
            message=classification.reason, tier=tier, quota=quota,
            classification=classification, claude_calls=calls,
        )
    if classification.status == "error":
        return AutofillRun(
            company_id=company_id, source_path=str(source), status=STATUS_ERROR,
            message=classification.reason, tier=tier, quota=quota,
            classification=classification, claude_calls=calls,
        )

    if not classification.proceed:
        message = (
            f"Skipped — this does not look like a tender document "
            f"({classification.confidence:.0%} confidence, {CONFIDENCE_THRESHOLD:.0%} needed). "
            f"{classification.reason}"
        )
        log.info("agent_autofill.orchestrator skipped below threshold: %s",
                 classification.summary_line)
        return AutofillRun(
            company_id=company_id, source_path=str(source),
            status=STATUS_SKIPPED_NOT_TENDER, message=message, tier=tier, quota=quota,
            classification=classification, claude_calls=calls,
        )

    # ---- 4. Extraction. -----------------------------------------------------
    try:
        report = extract_document(source)
    except Exception as exc:  # noqa: BLE001 - a malformed file is not a crash
        log.error("agent_autofill.orchestrator extraction failed on %s: %s", source.name, exc)
        return AutofillRun(
            company_id=company_id, source_path=str(source), status=STATUS_ERROR,
            message=f"Could not read the document structure: {exc}",
            tier=tier, quota=quota, classification=classification, claude_calls=calls,
        )

    # ---- 5. Fill. -----------------------------------------------------------
    #
    # PDFs are filled by writing at the blank's own coordinates (pdf_filler).
    # This used to return ANALYSIS_ONLY for anything that was not .docx, which
    # meant a real SA tender pack — three PDFs — produced nothing at all and
    # told the user to "ask the buyer for the Word version". Most buyers do not
    # have one. A .doc stays analysis-only: there is no pure-Python writer for
    # binary OLE2, and that is a real limit rather than a missing feature.
    profile = get_company_profile(company_id)
    profile, preference = _with_preference_claim(profile, report)
    out_root = Path(output_dir) if output_dir else file_paths.generated_dir()
    out_root.mkdir(parents=True, exist_ok=True)

    if report.doc_type not in ("docx", "pdf"):
        analysis = _pdf_analysis(report)
        run = AutofillRun(
            company_id=company_id, source_path=str(source), status=STATUS_ANALYSIS_ONLY,
            message=(f"This is a tender document ({classification.confidence:.0%} "
                     f"confidence), but it is a legacy .doc file. There is no way to "
                     f"write into one safely, so it was read for eligibility only."),
            tier=tier, quota=quota, classification=classification,
            extraction_summary=analysis, claude_calls=calls,
        )
        _consume_quota(run)
        return run

    is_pdf = report.doc_type == "pdf"
    draft_path = out_root / _output_name(source, ".pdf" if is_pdf else ".docx")

    try:
        if is_pdf:
            from agent_autofill.fill_engine.pdf_filler import fill_pdf

            fill_result = fill_pdf(source, draft_path, profile, match_label,
                                   company_id=company_id)
        else:
            fill_result = fill_docx(source, draft_path, profile, match_label)

        # A PDF with no detectable blanks is a notice or a specification, not a
        # form. Saying "0 filled" would imply a form we failed at; analysis-only
        # is the truthful description.
        if is_pdf and not fill_result.filled and not fill_result.skipped:
            analysis = _pdf_analysis(report)
            run = AutofillRun(
                company_id=company_id, source_path=str(source),
                status=STATUS_ANALYSIS_ONLY,
                message=(f"This is a tender document ({classification.confidence:.0%} "
                         f"confidence) with no fillable blanks — it reads as a notice "
                         f"or specification rather than a form. It was used for the "
                         f"eligibility check."),
                tier=tier, quota=quota, classification=classification,
                extraction_summary=analysis, claude_calls=calls,
            )
            _consume_quota(run)
            return run
    except Exception as exc:  # noqa: BLE001
        log.error("agent_autofill.orchestrator fill failed on %s: %s", source.name, exc)
        return AutofillRun(
            company_id=company_id, source_path=str(source), status=STATUS_ERROR,
            message=f"Could not write the draft: {exc}",
            tier=tier, quota=quota, classification=classification,
            extraction_summary=report.summary(), claude_calls=calls,
        )

    # ---- 6. Review summary. -------------------------------------------------
    summary_path = out_root / _output_name(source, ".html")
    summary_path.write_text(
        render_html(fill_result, company_name=str(profile.get("company_name") or "")),
        encoding="utf-8",
    )

    # ---- 7. STOP. --------------------------------------------------------
    # A draft and a summary exist. Nothing is reviewed, nothing is exportable.
    run = AutofillRun(
        company_id=company_id, source_path=str(source), status=STATUS_AWAITING_REVIEW,
        message=(f"{fill_result.summary_line}. Open the review summary, confirm every "
                 f"flagged field, and complete the signature, pricing and declaration "
                 f"sections yourself — Agent Autofill never fills those."),
        tier=tier, quota=quota, classification=classification,
        extraction_summary=report.summary(), fill_result=fill_result,
        output_document=str(draft_path), review_summary_path=str(summary_path),
        review_summary_url=file_paths.public_url(summary_path.name),
        claude_calls=calls,
    )
    _consume_quota(run)
    log.info("agent_autofill.orchestrator drafted %s -> %s (%s)",
             source.name, draft_path.name, fill_result.summary_line)
    return run


def _consume_quota(run: AutofillRun) -> None:
    """
    Burn one of the day's autofills — only after the work actually succeeded.

    Kept as its own function so there is exactly one place that writes to the
    usage log, and so the ordering rule is visible rather than implied.
    """
    if run.status not in _CONSUMES_QUOTA:
        return
    log_autofill_run(run.company_id)
    run.quota_consumed = True


def run_autofill_batch(
    company_id: str,
    paths,
    *,
    output_dir: str | Path | None = None,
) -> list[AutofillRun]:
    """
    Run a folder's worth of files, stopping the moment the quota is exhausted.

    The early break matters: without it, a 200-file folder on a Pro plan would
    pay for 200 classification calls to produce 3 drafts.
    """
    runs: list[AutofillRun] = []
    for path in paths:
        run = run_autofill(company_id, path, output_dir=output_dir)
        runs.append(run)
        if run.status in (STATUS_REFUSED_TIER, STATUS_REFUSED_QUOTA):
            break
    return runs
