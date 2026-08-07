"""Build the DOCX test form used by test_agent_autofill_extraction.py.

Replicates the SA MBD/SBD layout: two-column label/value tables, a merged-cell
row, a multi-column declaration grid (the MBD 4 "persons in service of the
state" table), and paragraph underscore blanks.

Run directly to (re)write tests/fixtures/mbd4_test_form.docx.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


# (label, prefilled_value) — "" means the cell is left blank to be completed.
SUPPLIER_INFORMATION: list[tuple[str, str]] = [
    ("NAME OF BIDDER", ""),
    ("Company Registration Number", ""),
    ("Tax Reference Number", ""),
    ("VAT Registration Number", ""),
    ("CSD Number", ""),
    ("B-BBEE Status Level", ""),
    ("POSTAL ADDRESS", ""),
    ("STREET ADDRESS", ""),
    ("CONTACT PERSON", ""),
    ("TELEPHONE NUMBER", ""),
    ("CELL PHONE NUMBER", ""),
    ("FACSIMILE NUMBER", ""),
    ("E-MAIL ADDRESS", ""),
    # A pre-filled row: the issuing department completes this, the bidder does
    # not. It must NOT be reported as a blank.
    ("BID NUMBER", "06/FY/27"),
    ("CLOSING DATE", "18 AUGUST 2026"),
]

PARTICULARS_OF_BIDDER: list[tuple[str, str]] = [
    ("Full Name of bidder or his or her representative", ""),
    ("Identity Number", ""),
    ("Position occupied in the company", ""),
    ("BIDDER NAME", ""),
    ("COMPANY REGISTRATION NUMBER", ""),
    ("BBBEE Status Level of Contribution", ""),
]


def build_mbd_test_form(path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading("MBD 1 - INVITATION TO BID", level=1)
    document.add_paragraph("SUPPLIER INFORMATION")

    # --- Table 0: classic two-column label/value grid -------------------
    t0 = document.add_table(rows=0, cols=2)
    t0.style = "Table Grid"
    for label, value in SUPPLIER_INFORMATION:
        row = t0.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    document.add_page_break()
    document.add_heading("MBD 4 - DECLARATION OF INTEREST", level=1)
    document.add_paragraph(
        "2.1 Any person having a controlling interest in the enterprise "
        "employed by the state must be declared below."
    )

    # --- Table 1: multi-column declaration grid (header row + empties) ---
    t1 = document.add_table(rows=1, cols=3)
    t1.style = "Table Grid"
    headers = ("Full Name", "Identity Number", "Name of State institution")
    for cell, text in zip(t1.rows[0].cells, headers):
        cell.text = text
    for _ in range(3):
        t1.add_row()  # three empty declaration rows

    document.add_page_break()
    document.add_heading("PARTICULARS OF BIDDER", level=1)

    # --- Table 2: label/value grid with a merged full-width note row -----
    t2 = document.add_table(rows=0, cols=2)
    t2.style = "Table Grid"
    for label, value in PARTICULARS_OF_BIDDER:
        row = t2.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    merged_row = t2.add_row()
    merged = merged_row.cells[0].merge(merged_row.cells[1])
    merged.text = ""  # one merged empty cell, must be reported exactly once

    # --- Paragraph underscore blanks ------------------------------------
    document.add_paragraph()
    document.add_paragraph("SIGNATURE: ______________________________")
    document.add_paragraph("CAPACITY: ______________________________")
    document.add_paragraph("DATE: ______________________________")
    document.add_paragraph(
        "This paragraph is ordinary prose with no blank in it at all."
    )

    document.save(str(path))
    return path


if __name__ == "__main__":
    out = Path(__file__).parent / "fixtures" / "mbd4_test_form.docx"
    build_mbd_test_form(out)
    print(f"wrote {out}")
