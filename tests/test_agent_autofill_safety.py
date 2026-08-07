"""
Adversarial tests for the Agent Autofill safety split.

The premise: a false positive on the blocklist costs a user one manual entry.
A false negative forges a signature or makes a false declaration to an organ of
state. These tests therefore try hard to get a never-fill field past the gate,
using phrasings that a pattern matcher tuned to "SIGNATURE" would miss.
"""

import os
import sys

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from agent_autofill.fill_engine.never_fill_fields import BlockReason, is_blocked
from agent_autofill.fill_engine.safe_fill_fields import decide, resolve_value

PROFILE = {
    "company_name": "CairoAI (Pty) Ltd",
    "registration_number": "2026/250499/07",
    "csd_number": "MAAA1234567",
    "bbbee_level": "Level 1 Contributor",
    "tax_reference_number": "9012345678",
    "vat_registration_number": "4480290011",
    "physical_address": "Centurion, Gauteng",
    "postal_address": "PO Box 1, Centurion",
    "standard_contact_person": "T. Molwantwa",
    "standard_phone": "+27 12 000 0000",
    "standard_email": "bids@cairoai.co.za",
    "directors": [{"name": "T. Molwantwa", "id_number": "9001015800086",
                   "is_state_employee": False}],
}


# --- signature: the adversarial set ---------------------------------------
# None of these contain the literal token "SIGNATURE:" that a naive
# implementation keys on.
SIGNATURE_VARIANTS = [
    "Sign here:",
    "Sign here",
    "Please sign below",
    "Signed at Ladysmith on this day",
    "SIGNED:",
    "Signatory",
    "Duly authorised signatory",
    "Handtekening",                      # Afrikaans, bilingual SA forms
    "Witness No. 1",
    "WITNESS 2",
    "Deponent",
    "Commissioner of Oaths",
    "Initials",
    "Initial each page",
    "Signature of Bidder",
    "SIGNATURE OF PERSON AUTHORISED",
    "sign off",
]


@pytest.mark.parametrize("label", SIGNATURE_VARIANTS)
def test_signature_variants_are_all_blocked(label):
    d = is_blocked(label)
    assert d.blocked, f"UNBLOCKED SIGNATURE LABEL: {label!r}"
    assert d.reason is BlockReason.SIGNATURE, f"{label!r} blocked but as {d.reason}"


@pytest.mark.parametrize("label", SIGNATURE_VARIANTS)
def test_signature_variants_never_produce_a_value(label):
    """Even with a whitelisted canonical field attached, nothing is written."""
    dec = decide("company_name", label, PROFILE, match_score=100.0)
    assert dec.fill is False
    assert dec.value is None


# --- declaration of interest ----------------------------------------------
DECLARATION_VARIANTS = [
    "Are you or any of your directors employed by the state?",
    "MBD 4",
    "SBD 4 Declaration of Interest",
    "Is any person connected to the bidder in the service of the state?",
    "Do you have a relationship with any person employed by the procuring institution?",
    "Public office bearer",
    "Any conflict of interest to declare",
    "family member employed by the municipality",
]


@pytest.mark.parametrize("label", DECLARATION_VARIANTS)
def test_declaration_variants_are_blocked(label):
    d = is_blocked(label)
    assert d.blocked, f"UNBLOCKED DECLARATION LABEL: {label!r}"
    assert d.reason is BlockReason.DECLARATION


def test_declaration_is_blocked_even_when_profile_holds_a_stored_answer():
    """
    The whole point: a stored 'no' from a previous tender must not be reused.
    The declaration is made afresh per tender, per date.
    """
    profile = dict(PROFILE)
    profile["directors"] = [{"name": "T. Molwantwa", "id_number": "9001015800086",
                             "is_state_employee": False}]
    dec = decide("director_names_and_id_numbers",
                 "Are any of the directors employed by the state?",
                 profile, match_score=100.0)
    assert dec.fill is False
    assert dec.block is not None and dec.block.reason is BlockReason.DECLARATION


# --- pricing ---------------------------------------------------------------
@pytest.mark.parametrize("label", [
    "Bid Amount", "Total price (incl VAT)", "Contract value", "Rate per hour",
    "Tariff", "Quotation total", "R ________", "Cost of works", "Fee proposal",
])
def test_pricing_is_never_written_here(label):
    d = is_blocked(label)
    assert d.blocked and d.reason is BlockReason.PRICING


# --- narrative -------------------------------------------------------------
@pytest.mark.parametrize("label", [
    "Method statement", "Describe your approach", "Motivation",
    "Please explain how you will deliver", "Project plan", "Reasons for",
])
def test_narrative_is_never_inserted(label):
    assert is_blocked(label).blocked


# --- signing dates vs factual dates ---------------------------------------
@pytest.mark.parametrize("label", [
    "Date of signature", "Signed on", "Thus done and signed",
    "Dated at", "this ___ day of",
])
def test_signing_dates_blocked(label):
    assert is_blocked(label).blocked


# --- the fields that SHOULD fill ------------------------------------------
@pytest.mark.parametrize("canonical,label,expected", [
    ("company_name", "NAME OF BIDDER", "CairoAI (Pty) Ltd"),
    ("registration_number", "Company Registration Number", "2026/250499/07"),
    ("csd_number", "CSD Number", "MAAA1234567"),
    ("tax_reference_number", "Tax Reference Number", "9012345678"),
    ("vat_registration_number", "VAT Registration Number", "4480290011"),
    ("physical_address", "Physical Address", "Centurion, Gauteng"),
    ("email_address", "E-mail address", "bids@cairoai.co.za"),
])
def test_safe_fields_do_fill(canonical, label, expected):
    dec = decide(canonical, label, PROFILE, match_score=100.0)
    assert dec.fill is True, dec.reason
    assert dec.value == expected


def test_directors_render_as_factual_roster():
    v, _ = resolve_value("director_names_and_id_numbers", PROFILE)
    assert v == "T. Molwantwa (9001015800086)"


# --- failure modes ---------------------------------------------------------
def test_unreadable_label_is_blocked_not_filled():
    """If extraction could not read the label, we must not write into it."""
    for empty in (None, "", "   "):
        assert is_blocked(empty).blocked


def test_low_score_match_is_not_filled():
    dec = decide("company_name", "Naam van bieër", PROFILE, match_score=70.0)
    assert dec.fill is False
    assert dec.low_confidence is True


def test_field_absent_from_profile_is_not_invented():
    thin = {"company_name": "CairoAI (Pty) Ltd"}
    dec = decide("vat_registration_number", "VAT Registration Number", thin, 100.0)
    assert dec.fill is False
    assert dec.value is None


def test_non_whitelisted_field_is_refused_even_with_perfect_match():
    dec = decide("capacity", "Capacity", PROFILE, match_score=100.0)
    assert dec.fill is False


def test_blocklist_beats_whitelist_on_conflict():
    """
    A label that maps to a safe field but reads as a signature line must be
    refused. This is the trap: 'Signature of the person whose name appears
    above' is name-adjacent.
    """
    dec = decide("company_name",
                 "Signature of the person whose name appears above",
                 PROFILE, match_score=100.0)
    assert dec.fill is False
    assert dec.block.reason is BlockReason.SIGNATURE


def test_tender_document_cannot_supply_values():
    """
    Values are drawn only from the confirmed profile. A counterparty document
    must never be able to change what we believe about our own company.
    """
    v, source = resolve_value("company_name", {"company_name": "CairoAI (Pty) Ltd"})
    assert v == "CairoAI (Pty) Ltd"
    assert source == "company profile"


# --- document-level context: found against the REAL SBD 4 -----------------
# The revised National Treasury form is headed "BIDDER'S DISCLOSURE", not
# "Declaration of Interest". Its declaration table is a grid of innocuous cells
# ("Full Name", "Identity Number", "Name of organ of state") that no per-label
# rule catches. Before the context guard, all 43 were merely "skipped" — safe
# only because the alias dictionary happens not to map "Full Name". Adding that
# obvious alias would have auto-populated a state-employee declaration.
from agent_autofill.fill_engine.never_fill_fields import classify_document_context

REAL_SBD4_TITLES = [
    "BIDDER'S DISCLOSURE",
    "BIDDER\u2019S DISCLOSURE",          # curly apostrophe, as in the real file
    "DECLARATION ON EMPLOYMENT BY ORGAN OF STATE",
    "GENERAL DECLARATION",
    "Declaration of Interest",             # legacy MBD 4 wording
    "SBD 4",
    "MBD 4",
]


@pytest.mark.parametrize("title", REAL_SBD4_TITLES)
def test_declaration_document_is_recognised_by_any_of_its_titles(title):
    assert "declaration_of_interest" in classify_document_context(title), title


@pytest.mark.parametrize("label", [
    "Full Name", "Identity Number", "ID Number", "Name of organ of state",
    "Supplier registration number (MAAA)", "Status (active/inactive/deleted)",
    "Position held",
])
def test_person_cells_are_blocked_inside_a_declaration_document(label):
    ctx = {"declaration_of_interest"}
    d = is_blocked(label, None, ctx)
    assert d.blocked, f"UNBLOCKED inside declaration form: {label!r}"
    assert d.reason is BlockReason.DECLARATION


@pytest.mark.parametrize("label", ["Full Name", "Identity Number"])
def test_same_labels_are_not_blocked_outside_a_declaration(label):
    """The guard must be context-specific, not a blanket ban on name fields."""
    assert not is_blocked(label, None, set()).blocked


def test_declaration_context_beats_a_whitelisted_canonical_field():
    """
    The scenario that motivated this: someone adds a plausible alias mapping
    'Full Name' to a fillable field. Inside an SBD 4 it must still refuse.
    """
    dec = decide("company_name", "Full Name", PROFILE, 100.0,
                 context={"declaration_of_interest"})
    assert dec.fill is False
    assert dec.block.reason is BlockReason.DECLARATION


# --- over-blocking found against the real MBD 1 supplier block -------------
# Both of these were blocked before the exemption rules: "CAPACITY UNDER WHICH
# THIS BID IS SIGNED" as a signature (the word SIGNED), and "TOTAL NUMBER OF
# ITEMS OFFERED" as pricing (a greedy \btotal\b). Both are plain facts.
@pytest.mark.parametrize("label", [
    "CAPACITY UNDER WHICH THIS BID IS SIGNED",
    "Capacity under which this bid is signed",
    "Capacity in which this bid is signed",
    "Designation",
    "Position held",
])
def test_capacity_is_a_fact_not_a_signature(label):
    assert not is_blocked(label).blocked, f"over-blocked: {label!r}"


@pytest.mark.parametrize("label", [
    "TOTAL NUMBER OF ITEMS OFFERED",
    "Total number of items",
    "Number of items offered",
    "Total quantity",
])
def test_item_counts_are_not_pricing(label):
    assert not is_blocked(label).blocked, f"over-blocked: {label!r}"


@pytest.mark.parametrize("label", [
    "Total price", "Total amount due", "Total (incl VAT)", "Contract total",
])
def test_money_totals_are_still_blocked(label):
    """The exemption must not open a hole for actual money fields."""
    assert is_blocked(label).blocked, f"UNBLOCKED MONEY FIELD: {label!r}"


@pytest.mark.parametrize("label", [
    "Signature", "Signature of Bidder", "SIGNED:", "Sign here:",
])
def test_signature_still_blocked_after_exemptions(label):
    assert is_blocked(label).blocked, f"UNBLOCKED SIGNATURE: {label!r}"


def test_exemption_never_overrides_declaration_context():
    """Inside an SBD 4, an exempt-looking label is still part of the declaration."""
    d = is_blocked("Position held", None, {"declaration_of_interest"})
    assert d.blocked and d.reason is BlockReason.DECLARATION


def test_landline_and_mobile_no_longer_collapse():
    """MBD 1 asks for both on adjacent rows; one value on both is wrong."""
    profile = dict(PROFILE)
    profile["standard_phone"] = "+27 12 000 0000"
    profile["standard_cell"] = "+27 82 555 1234"
    tel = decide("telephone_number", "TELEPHONE NUMBER", profile, 100.0)
    cell = decide("cell_phone_number", "CELLPHONE NUMBER", profile, 100.0)
    assert tel.value == "+27 12 000 0000"
    assert cell.value == "+27 82 555 1234"
    assert tel.value != cell.value


def test_mobile_absent_does_not_fall_back_to_landline():
    profile = dict(PROFILE)
    profile["standard_cell"] = None
    assert decide("cell_phone_number", "CELLPHONE NUMBER", profile, 100.0).fill is False


# --- table shapes ----------------------------------------------------------
# SA bid forms use two layouts and the filler must see both. Missing the second
# meant SBD 4's 43 declaration cells were neither filled NOR marked — safe in
# outcome, but the spec requires nothing be skipped without a visible trace.
import docx as _docx

from agent_autofill.extraction import match_label
from agent_autofill.fill_engine.document_filler import (
    SKIP_MARKER, _is_column_header_table, fill_docx,
)


def _label_value_doc(tmp_path):
    d = _docx.Document()
    t = d.add_table(rows=3, cols=2)
    for i, lbl in enumerate(["NAME OF BIDDER", "E-MAIL ADDRESS", "SIGNATURE OF BIDDER"]):
        t.rows[i].cells[0].text = lbl
        t.rows[i].cells[1].text = ""
    p = tmp_path / "label_value.docx"
    d.save(str(p))
    return p


def _column_header_doc(tmp_path):
    d = _docx.Document()
    t = d.add_table(rows=3, cols=3)
    for c, h in enumerate(["Full Name", "Identity Number", "Name of organ of state"]):
        t.rows[0].cells[c].text = h
    d.add_paragraph("BIDDER'S DISCLOSURE")          # makes it a declaration form
    p = tmp_path / "column_header.docx"
    d.save(str(p))
    return p


def test_column_header_table_is_recognised(tmp_path):
    d = _docx.Document(str(_column_header_doc(tmp_path)))
    assert _is_column_header_table(d.tables[0]) is True


def test_label_value_table_is_not_treated_as_column_header(tmp_path):
    d = _docx.Document(str(_label_value_doc(tmp_path)))
    assert _is_column_header_table(d.tables[0]) is False


def test_column_header_cells_are_all_marked_not_silently_skipped(tmp_path):
    src = _column_header_doc(tmp_path)
    out = tmp_path / "filled.docx"
    r = fill_docx(src, out, PROFILE, match_label)
    # 2 entry rows x 3 columns
    assert r.fillable_total == 6
    assert len(r.filled) == 0, "nothing in a declaration table may be filled"
    assert all(s.category == "blocked" for s in r.skipped)
    d = _docx.Document(str(out))
    marks = sum(1 for row in d.tables[0].rows for c in row.cells
                if c.text.strip() == SKIP_MARKER)
    assert marks == 6, f"expected every entry cell marked, got {marks}"


def test_label_value_path_still_fills(tmp_path):
    """Regression: adding the column-header path must not break the common one."""
    src = _label_value_doc(tmp_path)
    out = tmp_path / "filled2.docx"
    r = fill_docx(src, out, PROFILE, match_label)
    filled = {f.label: f.value for f in r.filled}
    assert filled.get("NAME OF BIDDER") == "CairoAI (Pty) Ltd"
    assert "SIGNATURE OF BIDDER" in {s.label for s in r.skipped}


def test_source_document_is_never_modified(tmp_path):
    import hashlib
    src = _label_value_doc(tmp_path)
    before = hashlib.sha256(src.read_bytes()).hexdigest()
    fill_docx(src, tmp_path / "out.docx", PROFILE, match_label)
    assert hashlib.sha256(src.read_bytes()).hexdigest() == before


def test_existing_answers_are_never_overwritten(tmp_path):
    d = _docx.Document()
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "NAME OF BIDDER"
    t.rows[0].cells[1].text = "Someone Else (Pty) Ltd"
    src = tmp_path / "prefilled.docx"
    d.save(str(src))
    out = tmp_path / "out3.docx"
    r = fill_docx(src, out, PROFILE, match_label)
    assert r.fillable_total == 0
    assert _docx.Document(str(out)).tables[0].rows[0].cells[1].text == "Someone Else (Pty) Ltd"


# --- counterparty sections: found on the real 06FY27 MBD 1 page -----------
# CONTACT PERSON / TELEPHONE NUMBER / E-MAIL ADDRESS appear twice on that page:
# under "BIDDING PROCEDURE ENQUIRIES MAY BE DIRECTED TO" (the buyer's own staff)
# and under SUPPLIER INFORMATION (us). Identical labels; only the section
# heading tells them apart.
from agent_autofill.fill_engine.never_fill_fields import is_counterparty_section


@pytest.mark.parametrize("section", [
    "BIDDING PROCEDURE ENQUIRIES MAY BE DIRECTED TO",
    "DIRECTED TO",
    "TECHNICAL ENQUIRIES MAY BE DIRECTED TO:",
    "FOR ATTENTION OF",
    "DEPARTMENT",
])
def test_counterparty_sections_are_recognised(section):
    assert is_counterparty_section(section), section


@pytest.mark.parametrize("section", ["SUPPLIER INFORMATION", "PART A", "", None])
def test_bidder_sections_are_not_counterparty(section):
    assert not is_counterparty_section(section)


@pytest.mark.parametrize("label", ["CONTACT PERSON", "TELEPHONE NUMBER", "E-MAIL ADDRESS"])
def test_our_details_never_go_in_the_buyers_block(label):
    d = decide("contact_person", label, PROFILE, 100.0, section="DIRECTED TO")
    assert d.fill is False, f"would have written our details into the buyer's block: {label}"


@pytest.mark.parametrize("canonical,label", [
    ("contact_person", "CONTACT PERSON"),
    ("email_address", "E-MAIL ADDRESS"),
])
def test_same_labels_still_fill_under_supplier_information(canonical, label):
    d = decide(canonical, label, PROFILE, 100.0, section="SUPPLIER INFORMATION")
    assert d.fill is True, d.reason
