"""
The signature that binds an export to the review record it claims.

Two bypasses were demonstrated against the export gate and recorded in c5eb068
rather than quietly patched:

  1. Writing `acknowledged_at` straight into `agent_memory.db` made
     `export_reviewed()` succeed and produce a genuine REVIEWED file.
  2. Calling `stamp_docx()` with `flags_open=0` produced a valid REVIEWED file
     while the database still said DRAFT with flags outstanding.

These tests re-run both, and pin the residual weakness that remains.

DATABASE ISOLATION — as in test_agent_autofill_export_gate.py: the environment
is set before any `agent.*` import, because db_paths and file_paths read it at
import time.
"""

import os
import shutil
import sqlite3
import sys
import tempfile
from datetime import datetime
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

_TEST_DB_DIR = Path(tempfile.mkdtemp(prefix="cairo-bind-test-db-"))
_TEST_GEN_DIR = Path(tempfile.mkdtemp(prefix="cairo-bind-test-gen-"))
os.environ["AGENT_DB_DIR"] = str(_TEST_DB_DIR)
os.environ["AGENT_GENERATED_DIR"] = str(_TEST_GEN_DIR)
os.environ.setdefault("AUTOFILL_STAMP_SECRET", "test-only-not-a-real-secret")

from agent.db_paths import AGENT_MEMORY_DB                                # noqa: E402
from agent_autofill.extraction.field_alias_dictionary import match_label   # noqa: E402
from agent_autofill.fill_engine.document_filler import fill_docx          # noqa: E402
from agent_autofill.integration import export_metadata as em              # noqa: E402
from agent_autofill.integration import review_gate as rg                  # noqa: E402
from agent_autofill.integration import stamp_signing as ss                # noqa: E402

FORM = ROOT / "tests" / "fixtures" / "sa_forms_generated" / "mbd1_supplier_info.docx"
COMPANY = "bind_test_co"

PROFILE = {
    "company_name": "CairoAI Bind Test (Pty) Ltd",
    "registration_number": "2026/250499/07",
    "csd_number": "MAAA1234567",
    "tax_reference_number": "9012345678",
    "vat_registration_number": "4480290011",
    "physical_address": "Centurion, Gauteng",
    "standard_contact_person": "T. Molwantwa",
    "standard_email": "bids@cairoai.co.za",
}


@pytest.fixture
def review(tmp_path):
    draft = tmp_path / "draft.docx"
    result = fill_docx(FORM, draft, PROFILE, match_label)
    assert result.skipped, "fixture must produce at least one flagged field"
    opened = rg.open_review(COMPANY, result, company_name=PROFILE["company_name"])
    yield opened["review_id"], result, draft
    rg.delete_review(COMPANY, opened["review_id"])


def _stored_status(review_id):
    """The review's own DRAFT/REVIEWED state — not get_review()'s call status."""
    with sqlite3.connect(str(AGENT_MEMORY_DB)) as conn:
        return conn.execute(
            "SELECT status FROM autofill_review WHERE review_id = ?", (review_id,)
        ).fetchone()[0]


def _ack_all(review_id):
    for item in rg.get_review(COMPANY, review_id)["outstanding"]:
        rg.acknowledge_field(
            COMPANY, review_id, item["item_key"],
            f"I will complete {item['label']} by hand before submitting.",
        )
    _confirm_values(review_id)

def _confirm_values(review_id, company=None):
    """
    The other half of a human review: confirming the values that WERE filled.
    Acknowledging the flags alone no longer produces a reviewed export — a form
    with nothing flagged used to export as REVIEWED with no human involvement
    at all.
    """
    co = company or COMPANY
    shown = rg.filled_values(co, review_id)
    return rg.confirm_filled_values(
        co, review_id, [v["item_key"] for v in shown["values"]])



# --- bypass 1: forging the acknowledgement rows ----------------------------


def test_direct_db_write_cannot_produce_a_reviewed_export(review):
    """BYPASS 1. Setting acknowledged_at by hand used to be enough."""
    review_id, _, _ = review
    with sqlite3.connect(str(AGENT_MEMORY_DB)) as conn:
        forged = conn.execute(
            "UPDATE autofill_review_item SET acknowledged_at = ?,"
            " acknowledged_note = ? WHERE review_id = ?",
            (datetime.now().isoformat(timespec="seconds"), "forged", review_id),
        ).rowcount
    assert forged > 0, "the forgery itself must land, or the test proves nothing"

    out = rg.export_reviewed(COMPANY, review_id)
    assert out["status"] == "error", "BYPASS 1: forged rows produced a REVIEWED export"
    assert out.get("tamper_detected") is True
    assert _stored_status(review_id) == "DRAFT"


def test_editing_an_acknowledged_timestamp_invalidates_it(review):
    """
    A subtler forgery: acknowledge legitimately, then backdate. The MAC covers
    the timestamp, so the row stops verifying.
    """
    review_id, _, _ = review
    _ack_all(review_id)
    assert not rg._unverifiable_acknowledgements(review_id)

    with sqlite3.connect(str(AGENT_MEMORY_DB)) as conn:
        conn.execute(
            "UPDATE autofill_review_item SET acknowledged_at = ? WHERE review_id = ?",
            ("2020-01-01T00:00:00", review_id),
        )
    assert rg._unverifiable_acknowledgements(review_id), "backdating went undetected"


def test_editing_the_note_invalidates_it(review):
    """What the person said they checked is covered too."""
    review_id, _, _ = review
    _ack_all(review_id)
    with sqlite3.connect(str(AGENT_MEMORY_DB)) as conn:
        conn.execute(
            "UPDATE autofill_review_item SET acknowledged_note = ? WHERE review_id = ?",
            ("something the reviewer never wrote", review_id),
        )
    assert rg._unverifiable_acknowledgements(review_id), "note rewrite went undetected"


# --- bypass 2: fabricating the stamp ---------------------------------------


def test_stamp_docx_refuses_a_reviewed_stamp_with_no_signature(review, tmp_path):
    """BYPASS 2. flags_open=0 with the DB still showing open flags."""
    review_id, _, draft = review
    forged_file = tmp_path / "forged.docx"
    shutil.copy2(draft, forged_file)

    stamp = em.ReviewStamp(
        review_id=review_id, status=em.STATUS_REVIEWED,
        flags_total=11, flags_open=0, filled_count=10,
        company_name=PROFILE["company_name"], source_document=str(FORM),
        acknowledged=[("f1", "fabricated", "never happened")],
    )
    with pytest.raises(em.ReviewStateError, match="signature"):
        em.stamp_docx(forged_file, stamp)


def test_an_unreviewed_stamp_still_needs_no_signature(review, tmp_path):
    """
    The requirement is on the claim, not on stamping. An UNREVIEWED banner
    says "do not submit" — nothing is gained by making that harder to write,
    and export_draft() depends on it.
    """
    review_id, _, draft = review
    target = tmp_path / "unreviewed.docx"
    shutil.copy2(draft, target)
    em.stamp_docx(target, em.ReviewStamp(
        review_id=review_id, status=em.STATUS_UNREVIEWED,
        flags_total=11, flags_open=11, filled_count=10,
    ))
    assert em.read_review_state(target)["content_status"] == em.STATUS_UNREVIEWED


def test_a_made_up_signature_fails_verification(review, tmp_path):
    """
    RESIDUAL WEAKNESS, pinned deliberately.

    stamp_docx only requires that a signature is present — it cannot check one,
    because it has the file and not the review record. So a forger who passes
    mac="x" still gets a document whose banner reads REVIEWED. What they cannot
    get is one that verifies. verify_export is therefore the authority on
    whether an export is genuine, and anything user-facing must call it rather
    than trusting the banner.
    """
    review_id, _, draft = review
    forged_file = tmp_path / "bogus_mac.docx"
    shutil.copy2(draft, forged_file)
    em.stamp_docx(forged_file, em.ReviewStamp(
        review_id=review_id, status=em.STATUS_REVIEWED,
        flags_total=11, flags_open=0, filled_count=10,
        mac="not-a-real-signature",
    ))

    # The banner does read REVIEWED — this is the part that is not fixed.
    assert em.read_review_state(forged_file)["content_status"] == em.STATUS_REVIEWED
    # But it does not verify against the record.
    verdict = rg.verify_export(forged_file, COMPANY, review_id)
    assert verdict["mac_verified"] is False


# --- the legitimate path ---------------------------------------------------


def test_a_real_export_verifies(review):
    review_id, _, _ = review
    _ack_all(review_id)
    out = rg.export_reviewed(COMPANY, review_id)
    assert out["status"] == "success", out.get("message")

    verdict = rg.verify_export(out["export_path"], COMPANY, review_id)
    assert verdict["mac_verified"] is True
    assert verdict["content_status"] == em.STATUS_REVIEWED
    assert verdict["channels_agree"] is True


def test_verification_without_a_record_reports_unchecked_not_failed(review):
    """"Not checked" and "failed the check" are different findings."""
    review_id, _, _ = review
    _ack_all(review_id)
    out = rg.export_reviewed(COMPANY, review_id)
    assert rg.verify_export(out["export_path"])["mac_verified"] is None


def test_a_valid_export_stops_verifying_once_the_record_is_edited(review):
    """
    The stamp is signed over the acknowledgement timestamps. Editing them after
    the fact invalidates an export that was genuine when it was made — which is
    the point: the file and the record must keep agreeing.
    """
    review_id, _, _ = review
    _ack_all(review_id)
    out = rg.export_reviewed(COMPANY, review_id)
    assert rg.verify_export(out["export_path"], COMPANY, review_id)["mac_verified"] is True

    with sqlite3.connect(str(AGENT_MEMORY_DB)) as conn:
        conn.execute(
            "UPDATE autofill_review_item SET acknowledged_at = ?"
            " WHERE review_id = ? AND item_key = (SELECT MIN(item_key) FROM"
            " autofill_review_item WHERE review_id = ?)",
            ("2020-01-01T00:00:00", review_id, review_id),
        )
    assert rg.verify_export(out["export_path"], COMPANY, review_id)["mac_verified"] is False


def test_cross_tenant_verification_fails(review):
    review_id, _, _ = review
    _ack_all(review_id)
    out = rg.export_reviewed(COMPANY, review_id)
    verdict = rg.verify_export(out["export_path"], "someone_else_co", review_id)
    assert verdict["mac_verified"] is False


# --- fail-closed -----------------------------------------------------------


def test_signing_fails_closed_when_the_secret_is_absent(monkeypatch):
    """
    No secret means no export, never an export without tamper-evidence.
    """
    monkeypatch.delenv(ss.SECRET_ENV_VAR, raising=False)
    assert ss.secret_available() is False
    with pytest.raises(ss.StampSecretMissing):
        ss.ack_mac("r1", "F01", "2026-08-08T10:00:00", "note")
    with pytest.raises(ss.StampSecretMissing):
        ss.matches("anything", ss.ack_payload("r1", "F01", "x", "y"))


def test_acknowledgement_refuses_without_the_secret(review, monkeypatch):
    review_id, _, _ = review
    item = rg.get_review(COMPANY, review_id)["outstanding"][0]
    monkeypatch.delenv(ss.SECRET_ENV_VAR, raising=False)
    with pytest.raises(ss.StampSecretMissing):
        rg.acknowledge_field(COMPANY, review_id, item["item_key"],
                             "I will complete this field by hand.")
    # And nothing was recorded.
    monkeypatch.setenv(ss.SECRET_ENV_VAR, "test-only-not-a-real-secret")
    assert rg.get_review(COMPANY, review_id)["outstanding_count"] > 0


# --- the MAC itself --------------------------------------------------------


def test_a_different_secret_produces_a_different_mac(monkeypatch):
    payload = ss.ack_payload("r1", "F01", "2026-08-08T10:00:00", "note")
    monkeypatch.setenv(ss.SECRET_ENV_VAR, "secret-a")
    a = ss.sign(payload)
    monkeypatch.setenv(ss.SECRET_ENV_VAR, "secret-b")
    assert ss.sign(payload) != a


def test_key_order_does_not_change_the_mac():
    """Canonical serialisation — otherwise verification is flaky, not strict."""
    p1 = {"v": 1, "kind": "ack", "review_id": "r", "item_key": "F01"}
    p2 = {"item_key": "F01", "review_id": "r", "kind": "ack", "v": 1}
    assert ss.sign(p1) == ss.sign(p2)


def test_empty_mac_never_matches():
    for empty in (None, "", "   "):
        assert ss.matches(empty, ss.ack_payload("r", "F01", "t", "n")) is False


# --- the third bypass: the stamp was portable ------------------------------
#
# Found by attacking the fix above. The stamp MAC ties the stamp to the review
# RECORD, but nothing tied it to the document's CONTENT — and it structurally
# could not, because the MAC is written into the file and can only cover the
# digest of the draft as it stood beforehand. So a genuine export's body could
# be rewritten, or the whole stamp lifted onto an unrelated document, and
# verification still passed.


def _rewrite_company_name(path, replacement="SOMEONE ELSE ENTIRELY (Pty) Ltd"):
    import docx

    document = docx.Document(str(path))
    changed = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if PROFILE["company_name"] in run.text:
                            run.text = run.text.replace(
                                PROFILE["company_name"], replacement)
                            changed += 1
    document.save(str(path))
    return changed


def test_altering_a_genuine_exports_content_breaks_verification(review, tmp_path):
    review_id, _, _ = review
    _ack_all(review_id)
    out = rg.export_reviewed(COMPANY, review_id)
    assert rg.verify_export(out["export_path"], COMPANY, review_id)["mac_verified"] is True

    tampered = tmp_path / "tampered.docx"
    shutil.copy2(out["export_path"], tampered)
    assert _rewrite_company_name(tampered) > 0, "the tamper itself must land"

    verdict = rg.verify_export(tampered, COMPANY, review_id)
    assert verdict["content_verified"] is False
    assert verdict["mac_verified"] is False


def test_a_stamp_cannot_be_moved_onto_another_document(review, tmp_path):
    """
    The sharper form: take every core property from a genuine export and paste
    it onto a document that was never reviewed at all.
    """
    import docx

    review_id, _, _ = review
    _ack_all(review_id)
    out = rg.export_reviewed(COMPANY, review_id)

    other = tmp_path / "never_reviewed.docx"
    result = fill_docx(FORM, other, PROFILE, match_label)
    assert result.skipped

    src = docx.Document(out["export_path"]).core_properties
    dst = docx.Document(str(other))
    for field in ("content_status", "category", "subject", "comments",
                  "keywords", "last_modified_by"):
        setattr(dst.core_properties, field, getattr(src, field))
    dst.save(str(other))

    # It reads as REVIEWED — the stamp is genuine, just not this file's.
    assert em.read_review_state(other)["content_status"] == em.STATUS_REVIEWED
    assert rg.verify_export(other, COMPANY, review_id)["mac_verified"] is False


def test_editing_the_stored_digest_to_match_a_tampered_file_also_fails(review, tmp_path):
    """
    The obvious follow-up for an attacker with database access: tamper with the
    file, then update the recorded digest to match. The digest is itself signed,
    so it fails too.
    """
    from agent_autofill.integration.stamp_signing import file_sha256

    review_id, _, _ = review
    _ack_all(review_id)
    out = rg.export_reviewed(COMPANY, review_id)

    tampered = tmp_path / "tampered2.docx"
    shutil.copy2(out["export_path"], tampered)
    _rewrite_company_name(tampered)

    with sqlite3.connect(str(AGENT_MEMORY_DB)) as conn:
        conn.execute(
            "UPDATE autofill_review SET final_sha256 = ? WHERE review_id = ?",
            (file_sha256(tampered), review_id),
        )
    assert rg.verify_export(tampered, COMPANY, review_id)["mac_verified"] is False


def test_a_byte_identical_genuine_export_still_verifies(review, tmp_path):
    """A plain copy is the same document and must keep verifying."""
    review_id, _, _ = review
    _ack_all(review_id)
    out = rg.export_reviewed(COMPANY, review_id)
    copied = tmp_path / "copy.docx"
    shutil.copy2(out["export_path"], copied)
    assert rg.verify_export(copied, COMPANY, review_id)["mac_verified"] is True


# --- check 8: the question the whole build turns on ------------------------
#
# "Could this system ever mark a document submission-ready without a human
# confirming every auto-filled field?"
#
# It could, and it needed no attack at all. The gate required acknowledgement
# only of the fields it could NOT fill, so the values it DID write were never
# confirmed by anyone. On a form where everything filled cleanly, a REVIEWED
# export took zero human involvement.


def _all_safe_form(path):
    """A form on which every field is safely fillable — no signature, no price,
    no declaration, no narrative. A supplier detail sheet; entirely plausible."""
    import docx

    d = docx.Document()
    table = d.add_table(rows=4, cols=2)
    for i, label in enumerate(["NAME OF BIDDER", "COMPANY REGISTRATION NUMBER",
                               "CSD NUMBER", "VAT REGISTRATION NUMBER"]):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = ""
    d.save(str(path))
    return path


@pytest.fixture
def zero_flag_review(tmp_path):
    src = _all_safe_form(tmp_path / "supplier_details.docx")
    result = fill_docx(src, tmp_path / "zf_draft.docx", PROFILE, match_label)
    assert result.filled and not result.skipped, (
        "this fixture only proves anything if nothing is flagged")
    opened = rg.open_review(COMPANY, result, company_name=PROFILE["company_name"])
    yield opened["review_id"], result
    rg.delete_review(COMPANY, opened["review_id"])


def test_a_form_with_no_flags_cannot_export_reviewed_without_a_human(zero_flag_review):
    """THE finding. Nothing flagged used to mean nothing to confirm."""
    review_id, result = zero_flag_review
    out = rg.export_reviewed(COMPANY, review_id)
    assert out["status"] == "error", (
        "BLOCKING: a REVIEWED export was produced with zero human involvement")
    assert len(out["unconfirmed_values"]) == len(result.filled)
    assert _stored_status(review_id) == "DRAFT"


def test_the_same_form_exports_once_the_values_are_confirmed(zero_flag_review):
    review_id, result = zero_flag_review
    shown = rg.filled_values(COMPANY, review_id)
    assert shown["count"] == len(result.filled)
    # Every value is available to show the user, with its label.
    assert all(v["label"] and v["value"] for v in shown["values"])

    rg.confirm_filled_values(
        COMPANY, review_id, [v["item_key"] for v in shown["values"]])
    out = rg.export_reviewed(COMPANY, review_id)
    assert out["status"] == "success"
    assert rg.verify_export(out["export_path"], COMPANY, review_id)["mac_verified"]


def test_a_partial_confirmation_is_refused_not_intersected(zero_flag_review):
    """Confirming some of what you were shown is not confirming it."""
    review_id, _ = zero_flag_review
    keys = [v["item_key"] for v in rg.filled_values(COMPANY, review_id)["values"]]
    out = rg.confirm_filled_values(COMPANY, review_id, keys[:-1])
    assert out["status"] == "error"
    assert out["missing"] == [keys[-1]]
    assert rg.export_reviewed(COMPANY, review_id)["status"] == "error"


def test_confirming_a_key_that_is_not_in_the_document_is_refused(zero_flag_review):
    review_id, _ = zero_flag_review
    keys = [v["item_key"] for v in rg.filled_values(COMPANY, review_id)["values"]]
    out = rg.confirm_filled_values(COMPANY, review_id, keys + ["V99"])
    assert out["status"] == "error"
    assert out["unknown"] == ["V99"]


def test_editing_a_value_after_confirmation_invalidates_it(zero_flag_review):
    """The confirmation covers what the person saw, not merely that they clicked."""
    review_id, _ = zero_flag_review
    keys = [v["item_key"] for v in rg.filled_values(COMPANY, review_id)["values"]]
    rg.confirm_filled_values(COMPANY, review_id, keys)
    assert not rg._values_unconfirmed(COMPANY, review_id)

    with sqlite3.connect(str(AGENT_MEMORY_DB)) as conn:
        conn.execute(
            "UPDATE autofill_review_fill SET value = ? WHERE review_id = ?"
            " AND item_key = ?", ("TAMPERED (Pty) Ltd", review_id, keys[0]))
    assert rg._values_unconfirmed(COMPANY, review_id)
    assert rg.export_reviewed(COMPANY, review_id)["status"] == "error"


def test_the_stamp_no_longer_claims_a_review_that_did_not_happen(zero_flag_review):
    """
    The old wording said "All 0 flagged field(s) were acknowledged by a person"
    on a document nobody had looked at.
    """
    review_id, result = zero_flag_review
    keys = [v["item_key"] for v in rg.filled_values(COMPANY, review_id)["values"]]
    rg.confirm_filled_values(COMPANY, review_id, keys)
    out = rg.export_reviewed(COMPANY, review_id)

    state = em.read_review_state(out["export_path"])
    assert "All 0 flagged" not in state["comments"]
    assert f"{len(result.filled)} pre-filled value(s) confirmed" in state["comments"]


def test_every_new_tool_is_reachable_by_the_model():
    """
    The quotation gate shipped as a dead end because nothing could reach the
    resolve step. Not repeating that.
    """
    from agent.tool_dispatch import TOOL_REGISTRY
    from agent_autofill.integration.autofill_tools import autofill_tools

    names = {t["name"] for t in autofill_tools}
    for tool in ("autofill_show_filled_values", "autofill_confirm_filled_values"):
        assert tool in names, f"{tool} has no schema"
        assert tool in TOOL_REGISTRY, f"{tool} is not dispatchable"
